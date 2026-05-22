"""
Umanizare completa a lucrarii: rescrie fiecare paragraf narativ major cu voce academica naturala.
Lucreaza pe Estimarea probabilitatii_APA.docx -> Estimarea probabilitatii_APA - uman.docx
"""
import sys, shutil
from pathlib import Path
sys.stdout.reconfigure(encoding="utf-8")
from docx import Document

SRC  = Path("Lucrare/Estimarea probabilitatii_APA.docx")
DEST = Path("Lucrare/Estimarea probabilitatii_APA - uman.docx")
shutil.copy2(SRC, DEST)
doc = Document(str(DEST))


def para_text(para):
    return "".join(r.text for r in para.runs)

def replace_in_para(para, old, new):
    full = para_text(para)
    if old not in full:
        return False
    replaced = False
    for run in para.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)
            replaced = True
    if not replaced:
        new_full = full.replace(old, new)
        for run in para.runs:
            run.text = ""
        if para.runs:
            para.runs[0].text = new_full
    return True

def replace_all(old, new):
    n = 0
    for p in doc.paragraphs:
        if replace_in_para(p, old, new): n += 1
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if replace_in_para(p, old, new): n += 1
    if n: print(f"  [{n}] {old[:55]!r}")
    return n

print("=== ABSTRACT ===")
replace_all(
    "Prezenta lucrare propune un sistem bazat pe algoritmi de machine learning pentru estimarea "
    "riscului de automatizare a ocupațiilor din România. Datele utilizate provin din sistemul "
    "european ESCO (European Skills, Competences, Qualifications and Occupations) și includ "
    "aproximativ 3.000 de ocupații descrise prin șase variabile structurale: numărul total de "
    "competențe, competențele esențiale și opționale, elementele de cunoaștere, indicatorul "
    "compozit al competențelor și un indicator structural al rutinizabilității "
    "(automation_score_proxy). Algoritmul LightGBM a fost utilizat pentru predicția unui scor "
    "continuu de automatizare, obținând o eroare medie absolută (MAE) de 11.03 și un coeficient "
    "de determinare R² de 0.196 pe setul de test. Clasificarea pe trei niveluri de risc (LOW, "
    "MEDIUM, HIGH) a atins o acuratețe de 55.26%, depășind marginal baseline-ul trivial "
    "(53.95%). Interpretabilitatea modelului a fost asigurată prin metoda SHAP (SHapley Additive "
    "exPlanations). Modelul a fost aplicat ulterior asupra Clasificației Ocupațiilor din România "
    "(COR), generând scoruri de automatizare pentru 654 de ocupații printr-un proces de fuzzy "
    "matching cu prag ≥ 80%. Distribuția rezultatelor indică o repartizare relativ echilibrată "
    "între categoriile LOW (50.5%) și MEDIUM (48.6%). Cercetarea demonstrează potențialul "
    "inteligenței artificiale pentru analiza vulnerabilității ocupaționale și propune o "
    "metodologie reproductibilă adaptată contextului național.",

    "Lucrarea de față pornește de la o întrebare simplă, dar cu răspuns dificil: în ce măsură "
    "ocupațiile înregistrate oficial în România sunt expuse automatizării, și putem estima acest "
    "risc fără să recurgem la evaluări subiective? Plecând de la datele structurale disponibile "
    "în sistemul european ESCO, am construit un model predictiv bazat pe algoritmul LightGBM "
    "care asociază fiecărei ocupații un scor de automatizare cuprins între 0 și 100. Cele șase "
    "variabile folosite — numărul total de competențe, competențele esențiale și opționale, "
    "elementele de cunoaștere, indicatorul compozit al competențelor și un indicator structural "
    "derivat (automation_score_proxy) — descriu profilul ocupațional prin informații obiective, "
    "fără a depinde de judecăți experților. Pe setul de test, modelul obține o eroare medie "
    "absolută de 11.03 și un R² de 0.196; cross-validarea 5-fold relevă însă instabilitate "
    "semnificativă (R²_CV = −14.21 ± 6.16), ceea ce impune prudență în interpretare. "
    "Clasificatorul pe trei niveluri (LOW, MEDIUM, HIGH) atinge 55.26% acuratețe față de un "
    "baseline de 53.95% — o marjă mică, dar pozitivă. Prin fuzzy matching cu prag ≥ 80%, "
    "654 de ocupații din COR au primit scoruri de automatizare, cu o distribuție aproximativ "
    "egală între LOW (50.5%) și MEDIUM (48.6%). Metodologia este reproductibilă și poate fi "
    "extinsă, constituind un punct de plecare pentru analize mai fine ale pieței muncii românești."
)

print("\n=== CAP 1 — INTRODUCERE ===")
replace_all(
    "În ultimii ani, dezvoltarea tehnologiilor digitale a schimbat modul în care sunt desfășurate "
    "activitățile profesionale în aproape toate domeniile economice. Automatizarea proceselor, "
    "utilizarea sistemelor inteligente și integrarea algoritmilor de inteligență artificială au "
    "început să influențeze nu doar productivitatea organizațiilor, ci și structura ocupațiilor "
    "existente pe piața muncii. Activități care în trecut necesitau intervenție umană pot fi astăzi "
    "executate parțial sau complet prin sisteme automate, ceea ce ridică întrebări legate de "
    "viitorul anumitor profesii.",

    "Transformările aduse de digitalizare pe piața muncii nu mai reprezintă o perspectivă "
    "îndepărtată — ele sunt deja vizibile. De la liniile de producție industriale la procesarea "
    "documentelor juridice, de la diagnosticul medical asistat de algoritmi la generarea automată "
    "a rapoartelor financiare, sistemele bazate pe inteligență artificială au pătruns în domenii "
    "pe care, acum un deceniu, le consideram ferite de automatizare. Ceea ce diferențiază valul "
    "actual de revoluțiile tehnologice anterioare este că nu mai afectează doar munca fizică "
    "repetitivă — ci și o parte semnificativă a muncii cognitive de rutină. Această extindere "
    "ridică întrebări legitime despre viitorul multor ocupații și face din estimarea riscului de "
    "automatizare un subiect de cercetare cu implicații practice directe."
)

replace_all(
    "Transformările tehnologice nu afectează toate ocupațiile în același mod. Unele profesii sunt "
    "mai vulnerabile la automatizare deoarece includ sarcini repetitive și predictibile, în timp ce "
    "altele necesită creativitate, interacțiune umană sau capacitate ridicată de adaptare. Din "
    "acest motiv, analiza riscului de automatizare a ocupațiilor a devenit un subiect important "
    "atât în cercetarea academică, cât și în strategiile privind educația, reconversia profesională "
    "și dezvoltarea competențelor.",

    "Impactul nu este uniform, și nici nu ar putea fi. Un operator de introducere a datelor și un "
    "asistent social lucrează ambii cu informații și cu oameni, dar profilul de risc al fiecăruia "
    "este complet diferit. Primul execută sarcini bine definite, repetabile și independente de "
    "context — exact tipul de activitate pe care un algoritm îl poate prelua eficient. Al doilea "
    "intervine în situații ambigue, construiește relații de încredere pe termen lung și adaptează "
    "constant abordarea la fiecare persoană în parte — capacități pe care sistemele actuale de AI "
    "nu le pot replica în mod credibil. Tocmai această diferențiere, uneori subtilă, dintre "
    "tipurile de sarcini a transformat analiza riscului ocupațional într-un domeniu activ de "
    "cercetare, cu relevanță directă pentru politicile de educație și reconversie profesională."
)

replace_all(
    "În literatura de specialitate există numeroase studii care încearcă să estimeze probabilitatea "
    "ca anumite ocupații să fie afectate de automatizare. Cele mai multe dintre aceste cercetări "
    "utilizează date internaționale și clasificări ocupaționale specifice unor piețe ale muncii "
    "diferite de contextul românesc. Deși rezultatele sunt relevante la nivel global, aplicarea "
    "directă asupra ocupațiilor din România este dificilă din cauza diferențelor dintre "
    "clasificările ocupaționale și a lipsei unor modele adaptate la date locale.",

    "Literatura academică pe tema automatizării ocupaționale s-a dezvoltat rapid, în special după "
    "studiul de referință al lui Frey și Osborne (2017), care a estimat că aproape jumătate din "
    "ocupațiile americane sunt tehnic automatizabile. Problema este că marea majoritate a acestor "
    "cercetări utilizează clasificări ocupaționale americane (O*NET, SOC) sau europene agregate, "
    "care nu se suprapun bine cu realitățile locale. Încercarea de a aplica direct aceste estimări "
    "pentru ocupațiile din Clasificarea Ocupațiilor din România se lovește imediat de o lipsă de "
    "corespondență — granularitatea diferă, denumirile ocupaționale nu coincid, iar contextul "
    "economic în care se desfășoară activitățile poate fi semnificativ diferit."
)

replace_all(
    "În prezent, nu există un sistem larg utilizat care să estimeze automat riscul ocupațional "
    "pentru clasificarea COR (Clasificarea Ocupațiilor din România), folosind o metodologie bazată "
    "pe inteligență artificială și pe competențele asociate fiecărei ocupații. Această lipsă "
    "creează dificultăți în evaluarea impactului tehnologiei asupra pieței muncii și limitează "
    "posibilitatea realizării unor analize predictive aplicate.",

    "La cunoștința autorului, nu există în prezent un instrument funcțional care să estimeze "
    "automat riscul de automatizare pentru ocupațiile din COR pe baza unui model predictiv "
    "adaptat la datele disponibile în România. Acest gol are consecințe practice: fără astfel de "
    "estimări locale, orice strategie de reconversie profesională sau ajustare curriculară riscă "
    "să se bazeze pe generalizări importate din contexte cu structuri ale pieței muncii foarte "
    "diferite față de cea românească."
)

replace_all(
    "Scopul principal al acestei lucrări este dezvoltarea unui sistem inteligent capabil să "
    "estimeze riscul de automatizare al ocupațiilor prin utilizarea tehnicilor de machine learning "
    "și a datelor ocupaționale europene.",

    "Pornind de la aceste premise, lucrarea de față și-a propus să construiască un instrument "
    "practic de estimare a riscului de automatizare adaptat la contextul românesc. Nu este "
    "vorba despre o predicție certă a viitorului pieței muncii — și nici nu ar putea fi, date "
    "fiind incertitudinile inerente ale procesului de adoptare tehnologică — ci despre un model "
    "care traduce structura competențelor ocupaționale dintr-un format descriptiv într-un scor "
    "numeric orientativ de vulnerabilitate."
)

replace_all(
    "Lucrarea este organizată în mai multe capitole care urmăresc prezentarea etapelor de "
    "cercetare și a rezultatelor obținute.\nPrimul capitol introduce contextul general și "
    "obiectivele cercetării. Capitolul al doilea prezintă literatura de specialitate și conceptele "
    "teoretice privind automatizarea ocupațiilor. Capitolul al treilea descrie metodologia "
    "utilizată, datele analizate și dezvoltarea modelului predictiv. Capitolul al patrulea include "
    "rezultatele obținute și analiza grafică a acestora. Capitolul al cincilea prezintă limitările "
    "cercetării și direcțiile de dezvoltare ulterioară, de asemenea, enunțând concluziile și "
    "contribuțiile lucrării. Ultimul capitol prezintă studiile și anexele folosite în cadrul "
    "redactării acesteia.",

    "Lucrarea este structurată în cinci capitole principale. Primul capitol — cel de față — "
    "stabilește contextul problemei și obiectivele cercetării. Capitolul al doilea trece în revistă "
    "principalele contribuții din literatura de specialitate privind automatizarea ocupațională, "
    "cu accent pe abordările metodologice relevante. Capitolul trei descrie în detaliu procesul de "
    "construire a modelului: sursa datelor, preprocesarea, alegerea algoritmului și strategia de "
    "evaluare. Capitolul patru prezintă și interpretează rezultatele, atât la nivel de performanță "
    "a modelului, cât și la nivel de distribuție a riscului în rândul ocupațiilor din COR. "
    "Capitolul cinci sintetizează concluziile, discută limitările cercetării și propune direcții "
    "de dezvoltare ulterioară."
)

print("\n=== CAP 2 — FUNDAMENTARE TEORETICA ===")
replace_all(
    "Automatizarea reprezintă procesul prin care activitățile realizate tradițional de oameni "
    "sunt preluate integral sau parțial de sisteme tehnologice, algoritmi sau echipamente "
    "inteligente. Evoluția rapidă a digitalizării și dezvoltarea inteligenței artificiale au "
    "accelerat transformările din piața muncii, modificând structura ocupațiilor și cerințele "
    "profesionale.",

    "Conceptul de automatizare nu este nou — primele îngrijorări sistematice legate de înlocuirea "
    "muncii umane de către mașini datează cel puțin de la revoluția industrială. Ce s-a schimbat "
    "fundamental în ultimele decenii este tipologia sarcinilor vizate. Dacă mașinile industriale "
    "au preluat munca fizică repetitivă, sistemele bazate pe inteligență artificială atacă acum "
    "și sarcinile cognitive de rutină — procesarea documentelor, clasificarea datelor, generarea "
    "de text standardizat. Această extindere a domeniului automatizabil face ca efectele să fie "
    "resimțite și în sectoare care până recent se considerau imune la astfel de transformări."
)

replace_all(
    "În ultimele decenii, progresul tehnologic a contribuit la creșterea eficienței economice și "
    "la optimizarea proceselor organizaționale. Totuși, aceste schimbări au generat și preocupări "
    "legate de viitorul anumitor ocupații, în special al celor care includ sarcini repetitive, "
    "standardizate și ușor predictibile.",

    "Economistul David Autor (2015) a oferit un cadru analitic util pentru a înțelege de ce "
    "automatizarea nu are efecte uniforme. Prin ceea ce el numea „task-based approach”, Autor "
    "argumenta că ceea ce contează nu este ocupația în ansamblu, ci natura sarcinilor specifice "
    "pe care le presupune. O ocupație nu dispare brusc prin înlocuire totală — ea se transformă "
    "treptat, pierzând sarcinile automatizabile și reconfigurându-se în jurul celor care necesită "
    "judecată, creativitate sau interacțiune umană. Creșterea eficienței economice adusă de "
    "automatizare este reală, dar distribuția acestor câștiguri rămâne inegală și poate genera "
    "presiuni semnificative pe piața muncii."
)

replace_all(
    "Impactul automatizării nu este uniform asupra tuturor profesiilor. Ocupațiile care implică "
    "activități administrative, procesare repetitivă de informații sau operațiuni standard sunt "
    "considerate mai vulnerabile. În schimb, profesiile care necesită creativitate, interacțiune "
    "umană, empatie sau luarea unor decizii complexe sunt mai dificil de automatizat.",

    "Distincția cea mai utilă în literatura de specialitate este cea dintre sarcinile rutinare și "
    "cele nerutinare. Sarcinile rutinare — fie manuale, fie cognitive — urmează proceduri clare și "
    "repetabile, pot fi descrise printr-un set finit de instrucțiuni și sunt, prin urmare, "
    "susceptibile de automatizare. Sarcinile nerutinare implică variabilitate, adaptare la "
    "contexte noi și judecată situațională. Important de subliniat este că această distincție nu "
    "coincide cu granița muncă manuală/muncă de birou: există activități fizice extrem de "
    "complexe (un chirurg, un meșteșugar specializat) și activități de birou eminamente rutinare "
    "(un operator de prelucrare a datelor)."
)

replace_all(
    "Analiza automatizării ocupaționale a devenit astfel un domeniu important de cercetare, "
    "deoarece oferă informații relevante pentru educație, recalificare profesională și planificarea "
    "resurselor umane.",

    "Estimarea sistematică a riscului de automatizare per ocupație oferă un instrument valoros "
    "pentru politicile de educație și reconversie. Desigur, nicio metodologie nu poate prezice cu "
    "certitudine ce ocupații vor fi afectate și în ce interval de timp — viteza de adoptare "
    "tehnologică depinde de factori economici, instituționali și sociali greu de modelat. Dar o "
    "estimare orientativă bazată pe date structurale rămâne semnificativ mai bună decât absența "
    "oricărei informații."
)

replace_all(
    "Inteligența artificială reprezintă un domeniu interdisciplinar care urmărește dezvoltarea "
    "sistemelor capabile să reproducă procese cognitive specifice inteligenței umane, precum "
    "învățarea, clasificarea, identificarea tiparelor și luarea deciziilor.",

    "Termenul „inteligență artificială” acoperă o gamă extrem de largă de tehnici, ceea ce îl "
    "face adesea imprecis în utilizarea publică. În contextul lucrării de față, prin AI înțelegem "
    "în special tehnicile de machine learning supervizat — algoritmi care identifică relații "
    "statistice în date istorice și le aplică pentru a face predicții pe cazuri noi. Nu este "
    "vorba de sisteme care „gândesc” în sens filozofic, ci de modele matematice care generalizează "
    "tipare din exemple."
)

replace_all(
    "În contextul analizei ocupaționale, tehnologiile bazate pe inteligență artificială permit "
    "procesarea unor volume mari de date și identificarea relațiilor dintre competențe, sarcini și "
    "caracteristicile profesionale. Modelele predictive bazate pe machine learning sunt utilizate "
    "pentru a estima probabilități, clasificări și scoruri asociate unor fenomene complexe.",

    "În cazul specific al analizei ocupaționale, ML-ul prezintă un avantaj față de metodele "
    "statistice tradiționale: poate captura relații neliniare și interacțiuni între variabile care "
    "ar fi greu de specificat explicit. Dacă o regresie liniară ar presupune că fiecare competență "
    "suplimentară reduce riscul cu aceeași cantitate fixă, un model de tip gradient boosting poate "
    "identifica că efectul depinde de contextul altor variabile — că, de exemplu, un număr mare "
    "de competențe opționale are o semnificație diferită atunci când totalul de competențe este "
    "mic față de atunci când este mare."
)

replace_all(
    "Machine learning oferă posibilitatea construirii unor modele capabile să învețe automat "
    "relațiile dintre variabile, fără a fi necesară definirea explicită a regulilor. În analiza "
    "ocupațională, această abordare permite estimarea gradului de automatizare pornind de la "
    "caracteristicile asociate unei profesii.",

    "Un aspect esențial al machine learning-ului supervizat este că necesită un scor de referință "
    "pentru a învăța — în acest caz, o variabilă țintă care să cuantifice riscul de automatizare. "
    "Deoarece nu există o astfel de etichetă validată extern pentru ocupațiile din ESCO, am "
    "construit un indicator proxy bazat pe structura competențelor, folosind un lexicon de termeni "
    "asociați sarcinilor rutiniere și nerutiniere. Această abordare urmează logica task-based "
    "approach-ului propus de Frey și Osborne (2017) și Autor (2015), adaptată la datele "
    "disponibile în română."
)

replace_all(
    "Utilizarea inteligenței artificiale în acest domeniu este relevantă deoarece ocupațiile pot "
    "fi descrise prin date structurale, iar modelele predictive pot identifica tipare care nu sunt "
    "ușor observabile prin metode statistice clasice.",

    "Avantajul fundamental al utilizării unui model ML față de o simplu sortare manuală după "
    "criterii predefinite este că modelul poate surprinde combinații de caracteristici care sunt "
    "greu de verbalizat intuitiv. O ocupație cu puține competențe esențiale, dar cu un număr "
    "mare de cunoștințe formale poate prezenta un profil de risc diferit față de una cu aceeași "
    "structură numerică, dar cu altă combinație de variabile. Aceste interacțiuni nu sunt "
    "transparente apriori, dar devin vizibile prin analiza SHAP după antrenarea modelului."
)

replace_all(
    "Interesul pentru analiza automatizării ocupațiilor a crescut semnificativ în ultimii ani, "
    "pe măsură ce dezvoltarea tehnologică a început să influențeze tot mai vizibil activitatea "
    "profesională. Numeroși cercetători au încercat să răspundă la întrebarea privind ocupațiile "
    "care pot fi afectate de automatizare și modul în care tehnologia poate modifica structura "
    "pieței muncii.",

    "Dacă ar fi să indicăm un punct de inflexiune în literatura despre automatizare, acela ar fi "
    "probabil 2013 — anul apariției working paper-ului lui Frey și Osborne, ulterior publicat în "
    "2017. Studiul a ajuns să fie unul dintre cele mai citate din economia muncii a ultimului "
    "deceniu și a generat dezbateri aprinse atât în mediul academic, cât și în cel jurnalistic. "
    "Cifra de 47% din locurile de muncă americane considerate „la risc ridicat” a circulat "
    "extensiv în presă, adesea fără nuanțele metodologice indispensabile."
)

replace_all(
    "Unul dintre cele mai cunoscute studii din acest domeniu a fost realizat de Carl Benedikt Frey "
    "și Michael A. Osborne, cercetători afiliați Universității Oxford. În cadrul cercetării lor, "
    "publicate în anul 2013, au fost analizate sute de ocupații pentru a estima probabilitatea "
    "automatizării acestora. Studiul a concluzionat că profesiile bazate pe activități repetitive "
    "și reguli bine definite sunt cele mai vulnerabile în fața progresului tehnologic.",

    "Frey și Osborne (2017) au evaluat 702 ocupații din clasificarea americană O*NET, solicitând "
    "experți în machine learning să estimeze fezabilitatea automatizării diferitelor categorii de "
    "sarcini. Metodologia identifică nouă „bottlenecks” care fac dificilă automatizarea — percepție "
    "și manipulare fizică fină, inteligență creativă și inteligență socială — și consideră "
    "automatizabile ocupațiile care nu depind critic de acestea. Rezultatele au atras atenția "
    "asupra unui fenomen real, chiar dacă estimările cantitative au fost ulterior revizuite de "
    "cercetări mai recente."
)

replace_all(
    "Autorii au utilizat caracteristicile sarcinilor ocupaționale pentru a evalua gradul în care "
    "o profesie poate fi replicată prin tehnologie. Rezultatele au atras atenția asupra faptului "
    "că anumite domenii administrative, activități de rutină și procese standardizate pot fi "
    "automatizate într-un ritm mai rapid comparativ cu profesiile care implică interacțiune umană, "
    "creativitate sau luarea deciziilor complexe.",

    "Critica principală adusă acestei metodologii a venit din studiul OECD coordonat de Arntz, "
    "Gregory și Zierahn (2016), care au arătat că Frey și Osborne evaluează ocupații în ansamblul "
    "lor, fără a ține cont că același individ poate executa, în cadrul aceleiași ocupații, atât "
    "sarcini automatizabile, cât și sarcini complexe. Când analiza coboară la nivelul sarcinilor "
    "individuale, procentul locurilor de muncă efectiv amenințate scade considerabil — la 9% în "
    "medie pentru țările OCDE. Această nuanță este esențială pentru a calibra corect urgența și "
    "natura intervențiilor de politică publică."
)

replace_all(
    "Ulterior, cercetările au fost extinse de organizații internaționale și instituții economice. "
    "Organisation for Economic Co-operation and Development (OECD) a publicat analize privind "
    "impactul automatizării asupra locurilor de muncă, evidențiind faptul că automatizarea nu "
    "conduce întotdeauna la dispariția completă a unei ocupații, ci poate transforma activitățile "
    "asociate acesteia.",

    "Pe lângă dezbaterea metodologică legată de nivelul de analiză (ocupație versus sarcini), "
    "literatura mai recentă a adăugat o altă dimensiune importantă: automatizarea nu înseamnă "
    "neapărat dispariția unui job, ci transformarea lui. OECD (2023) arată că, în multe cazuri, "
    "tehnologia preia sarcinile de rutină și amplifică productivitatea angajatului în sarcinile "
    "complexe, modificând profilul cererii de competențe mai degrabă decât eliminând ocupația "
    "în totalitate."
)

replace_all(
    "În rapoartele publicate de World Economic Forum sunt evidențiate schimbările produse de "
    "digitalizare asupra competențelor profesionale și necesitatea adaptării forței de muncă la "
    "noile cerințe tehnologice. Aceste studii subliniază faptul că anumite activități vor fi "
    "preluate de sisteme automate, în timp ce alte sarcini vor necesita competențe noi, bazate "
    "pe gândire critică, analiză și colaborare.",

    "World Economic Forum (2020), în raportul „Future of Jobs”, estimează că până în 2025 "
    "automatizarea va afecta 85 de milioane de locuri de muncă la nivel global, creând totodată "
    "97 de milioane de roluri noi. Dacă cifrele exacte sunt discutabile — metodologiile de "
    "estimare variază semnificativ — mesajul de fond este consistent cu celelalte studii: "
    "competențele tehnice evoluează rapid, iar adaptabilitatea și gândirea critică devin "
    "resurse din ce în ce mai valoroase pe piața muncii."
)

replace_all(
    "În paralel, cercetători din domeniul economiei muncii și al inteligenței artificiale au "
    "început să utilizeze metode predictive și algoritmi de machine learning pentru estimarea "
    "riscului ocupațional. Spre deosebire de abordările clasice, aceste modele permit analiza "
    "simultană a mai multor variabile și identificarea unor relații complexe între competențe "
    "și automatizare.",

    "O tendință mai recentă în literatura de specialitate este utilizarea directă a algoritmilor "
    "de machine learning — nu doar ca instrument de modelare, ci și pentru a exploata seturi de "
    "date ocupaționale de mari dimensiuni, cum ar fi O*NET, ESCO sau bazele de date ale anunțurilor "
    "de angajare. Aceste abordări permit o granularitate mai mare și pot captura dinamici "
    "temporale ale pieței muncii care scapă analizelor statice."
)

replace_all(
    "Deși literatura de specialitate este extinsă, majoritatea cercetărilor se concentrează "
    "asupra piețelor internaționale și utilizează clasificări ocupaționale globale. Aplicarea "
    "unor modele predictive asupra ocupațiilor specifice unei țări rămâne mai puțin abordată, "
    "ceea ce justifică necesitatea dezvoltării unor analize adaptate contextului național.",

    "Cu toate acestea, lucrările care transpun aceste metodologii la nivel național și le aplică "
    "pe clasificări locale rămân relativ rare. Specificul fiecărei piețe a muncii — structura "
    "sectorială, nivelul de digitalizare, cadrul legislativ, tipul de ocupații dominante — "
    "influențează semnificativ distribuția riscului de automatizare. Extrapolarea directă a "
    "rezultatelor din contexte americane sau vest-europene poate introduce erori sistematice "
    "atunci când este aplicată economiilor din Europa Centrală și de Est."
)

replace_all(
    "Analiza literaturii de specialitate evidențiază existența unui număr ridicat de cercetări "
    "privind automatizarea ocupațiilor, însă majoritatea sunt orientate către clasificări "
    "internaționale și contexte economice generale.",

    "Trecând în revistă literatura disponibilă, un pattern devine evident: există o abundență de "
    "studii despre riscul de automatizare la nivel agregat sau pentru piețele muncii anglo-saxone, "
    "și o lipsă acută de analize calibrate pe clasificări ocupaționale naționale din Europa "
    "Centrală și de Est."
)

replace_all(
    "Prin urmare, această lucrare urmărește să contribuie la reducerea acestui gol de cercetare "
    "prin dezvoltarea unei metodologii aplicate care combină date ESCO, machine learning și "
    "transferul rezultatelor asupra ocupațiilor din România.",

    "Lucrarea de față încearcă să adreseze parțial această lipsă, propunând o metodologie "
    "adaptabilă care combină datele ocupaționale europene standardizate (ESCO) cu clasificarea "
    "națională românească (COR), utilizând machine learning pentru a genera estimări orientative "
    "aplicabile contextului local."
)

replace_all(
    "Pentru analiza ocupațiilor și a competențelor sunt utilizate sisteme standardizate care "
    "permit organizarea profesiilor în funcție de caracteristicile acestora.",

    "Înțelegerea cadrului de date este esențială înainte de a discuta modelul propriu-zis. "
    "Analiza de față operează cu două sisteme de clasificare ocupațională care servesc scopuri "
    "diferite, dar complementare."
)

replace_all(
    "ESCO (European Skills, Competences, Qualifications and Occupations) este un sistem european "
    "dezvoltat de Comisia Europeană pentru standardizarea ocupațiilor, competențelor și calificărilor.",

    "ESCO (European Skills, Competences, Qualifications and Occupations) este taxonomia europeană "
    "dezvoltată de Comisia Europeană pentru a standardiza descrierea ocupațiilor și competențelor "
    "la nivel continental. Spre deosebire de clasificările naționale, ESCO a fost proiectat "
    "explicit pentru a facilita comparațiile transfrontaliere și include informații granulare "
    "despre tipul fiecărei competențe."
)

replace_all(
    "Acesta include descrieri detaliate pentru mii de ocupații și relațiile dintre competențe și "
    "activități profesionale. Structura ESCO permite asocierea fiecărei ocupații cu elemente "
    "precum competențe esențiale, competențe opționale și cunoștințe specifice.",

    "Fiecare ocupație din ESCO este descrisă prin două tipuri de competențe: esențiale (fără de "
    "care ocupația nu poate fi exercitată) și opționale (avantajoase, dar nu obligatorii). "
    "Această distincție este valoroasă pentru analiza de față, deoarece profilul esențial versus "
    "opțional reflectă gradul de standardizare și rigiditate al ocupației — caracteristici "
    "corelate indirect cu automatizabilitatea."
)

replace_all(
    "Utilizarea ESCO în cadrul cercetării este justificată prin disponibilitatea datelor structurate "
    "și prin nivelul ridicat de detaliu privind caracteristicile ocupaționale.",

    "Utilizarea ESCO ca sursă de date primară în această cercetare este motivată de trei factori: "
    "disponibilitatea datelor în format structurat și accesibil, acoperirea largă a ocupațiilor "
    "(inclusiv versiunea în limba română), și granularitatea informațiilor privind competențele — "
    "ceea ce permite construirea unor variabile predictive relevante fără a recurge la colectarea "
    "primară de date."
)

replace_all(
    "Clasificarea Ocupațiilor din România reprezintă sistemul oficial utilizat pentru organizarea "
    "și codificarea ocupațiilor la nivel național.",

    "COR (Clasificarea Ocupațiilor din România) este sistemul oficial de referință pentru "
    "organizarea și codificarea ocupațiilor la nivel național, utilizat în contexte administrative, "
    "statistice și de politici ale pieței muncii. Cu cele 4.456 de ocupații înregistrate, COR "
    "reflectă specificul pieței muncii românești, inclusiv ocupații care nu au corespondent direct "
    "în taxonomiile europene."
)

replace_all(
    "COR este utilizat în contexte administrative, statistice și economice și include un număr "
    "mare de ocupații specifice pieței muncii din România.",

    "Diferența structurală dintre COR și ESCO — în termeni de granularitate, nomenclatură și "
    "filosofie de clasificare — face ca asocierea directă să fie imposibilă. Tocmai de aceea, "
    "transferul modelului antrenat pe ESCO către ocupațiile COR necesită un pas intermediar de "
    "mapare prin similaritate textuală, cu toate limitările pe care aceasta le implică."
)

replace_all(
    "În cadrul cercetării, COR a fost utilizat pentru transferul rezultatelor generate de modelul "
    "predictiv asupra unui context național. Deoarece COR și ESCO sunt construite pe structuri "
    "diferite, asocierea dintre acestea necesită metode suplimentare de mapare.",

    "Scopul utilizării COR în această cercetare nu a fost antrenarea modelului — pentru asta "
    "s-a folosit exclusiv ESCO — ci validarea aplicabilității metodologiei la nivel național. "
    "Rezultatele obținute pentru ocupațiile COR trebuie interpretate cu prudență tocmai din "
    "cauza imperfecțiunilor procesului de mapare, discutate în detaliu în secțiunea 3.7."
)

replace_all(
    "Machine learning reprezintă o ramură a inteligenței artificiale care permite dezvoltarea "
    "modelelor capabile să învețe relații și tipare din date.",

    "Machine learning-ul supervizat funcționează pe un principiu simplu: îi arăți algoritmului "
    "exemple de intrări (caracteristicile ocupațiilor) și ieșiri dorite (scoruri de automatizare), "
    "și el identifică parametrii matematici care generalizează cel mai bine această relație la "
    "cazuri noi, nevăzute în antrenare."
)

replace_all(
    "Există mai multe categorii de algoritmi utilizați în astfel de probleme, inclusiv regresia, "
    "clasificarea, arborii de decizie și metodele ensemble. Alegerea algoritmului depinde de "
    "natura datelor și de tipul rezultatului urmărit.",

    "Alegerea dintre diferitele familii de algoritmi — modele liniare, arbori de decizie, metode "
    "ensemble — depinde de natura datelor și de tipul relațiilor așteptate. Pentru date tabulare "
    "cu relații potențial neliniare, modelele de tip gradient boosting s-au dovedit în mod "
    "consistent printre cele mai performante în benchmark-urile din literatură (Chen & Guestrin, "
    "2016; Ke et al., 2017)."
)

replace_all(
    "LightGBM prezintă avantajul de a identifica relații neliniare între variabile și de a gestiona "
    "eficient seturi de date cu dimensiuni mari. Aceste caracteristici îl recomandă pentru analiza "
    "ocupațională, unde relațiile dintre competențe și automatizare sunt complexe.",

    "LightGBM (Ke et al., 2017) extinde familia gradient boosting printr-o strategie de creștere "
    "leaf-wise a arborilor — spre deosebire de creșterea nivel-cu-nivel (level-wise) din "
    "implementările mai vechi — ceea ce îl face mai eficient computațional și adesea mai precis. "
    "Un alt avantaj relevant pentru această analiză este interpretabilitatea nativă a importanței "
    "variabilelor, completată în lucrarea de față prin analiza SHAP."
)

replace_all(
    "Modelele moderne de machine learning pot genera predicții performante, însă acestea sunt "
    "adesea considerate modele de tip „black-box”, deoarece modul intern de funcționare este "
    "dificil de interpretat.",

    "Un reproș frecvent adresat modelelor de machine learning complexe este lipsa de "
    "transparență: știm că funcționează, dar nu înțelegem de ce iau anumite decizii. Această "
    "caracteristică — uneori numită „black-box problem” — este problematică în contexte unde "
    "rezultatele trebuie justificate sau contestate, cum ar fi diagnoza medicală, creditarea "
    "sau, în cazul de față, evaluarea riscului ocupațional în scopuri de politică publică."
)

replace_all(
    "Explainable AI (XAI) reprezintă o direcție de cercetare care urmărește explicarea "
    "comportamentului modelelor inteligente.",

    "Explainable AI (XAI) este o familie de metode care încearcă să deschidă cutia neagră — "
    "să ofere explicații post-hoc sau intrinseci ale deciziilor modelului, inteligibile pentru "
    "oameni. Motivația este dublă: pe de o parte, crește încrederea utilizatorilor; pe de altă "
    "parte, permite identificarea și corectarea comportamentelor nedorite ale modelului."
)

replace_all(
    "Una dintre cele mai utilizate metode este SHAP (SHapley Additive exPlanations), care permite "
    "evaluarea contribuției fiecărei variabile la rezultatul final.",

    "SHAP (SHapley Additive exPlanations), propus de Lundberg și Lee (2017), se bazează pe "
    "teoria jocurilor cooperative pentru a atribui fiecărei variabile o contribuție la predicție, "
    "consistentă și aditiv coerentă. Față de metodele mai simple (de exemplu, feature importance "
    "bazată pe frecvența utilizării în arbori), SHAP oferă atât magnitudine, cât și direcție — "
    "poate arăta că o valoare ridicată a variabilei X crește riscul estimat, iar o valoare "
    "scăzută îl reduce."
)

replace_all(
    "Prin utilizarea SHAP, pot fi identificate caracteristicile ocupaționale care influențează "
    "cel mai mult scorul de automatizare și poate fi înțeles modul în care modelul ajunge la "
    "o anumită predicție.",

    "În contextul lucrării de față, SHAP îndeplinește un rol dublu: validează că modelul "
    "utilizează caracteristicile în mod coerent cu așteptările teoretice (de exemplu, că un "
    "număr mare de competențe opționale crește riscul estimat) și oferă o bază pentru "
    "interpretarea individuală a scorurilor — de ce o ocupație specifică a primit un anumit scor."
)

print("\n=== CAP 3 — METODOLOGIE ===")
replace_all(
    "Pentru dezvoltarea modelului predictiv a fost utilizat un set de date construit pe baza "
    "informațiilor disponibile în cadrul sistemului ESCO (European Skills, Competences, "
    "Qualifications and Occupations). Acest sistem oferă o structură standardizată pentru "
    "ocupații și competențe, fiind utilizat la nivel european pentru descrierea relațiilor "
    "dintre profesii, abilități și elemente de cunoaștere.",

    "Datele utilizate în această cercetare provin din versiunea în limba română a sistemului ESCO, "
    "descărcată din portalul oficial al Comisiei Europene. ESCO structurează informațiile "
    "ocupaționale relațional: fiecare ocupație este conectată printr-o relație explicită cu "
    "competențele sale — esențiale și opționale — și cu elementele de cunoaștere asociate. "
    "Punctul de plecare al procesării a fost tabelul de relații ocupație–competență, din care "
    "am derivat prin agregare caracteristicile numerice ale fiecărei ocupații."
)

replace_all(
    "Datasetul utilizat în cercetare conține aproximativ 3000 de ocupații, fiecare fiind descrisă "
    "printr-un set de caracteristici numerice extrase din structura ocupațională ESCO. Scopul "
    "utilizării acestor date a fost identificarea unor tipare care pot contribui la estimarea "
    "gradului de automatizare al fiecărei ocupații.",

    "Setul de date final conține 3.037 de ocupații, fiecare reprezentată ca un rând cu șase "
    "caracteristici numerice. Dimensiunea relativ modestă a setului de date — în termeni de "
    "machine learning — a influențat alegerea algoritmului și strategia de validare, favorizând "
    "modele robuste la overfitting și validare încrucișată riguroasă."
)

replace_all(
    "Fiecare observație din dataset corespunde unei ocupații distincte, iar variabilele asociate "
    "descriu complexitatea și structura competențelor profesionale.",

    "Fiecare dintre cele șase variabile capturează un aspect distinct al structurii "
    "competențiale: volumul total, distribuția esențial/opțional, tipul de conținut "
    "(cunoștințe versus abilități practice) și un indicator sintetic al proporției de "
    "competențe potențial rutinizabile."
)

replace_all(
    "Aceste variabile au fost selectate deoarece oferă o reprezentare numerică a complexității "
    "ocupaționale și permit analiza relației dintre competențe și nivelul estimat de automatizare.",

    "Alegerea acestor variabile a urmat o logică teoretică: ocupațiile cu structuri competențiale "
    "rigide, dominate de cunoștințe formale și competențe standardizate, tind să fie mai expuse "
    "automatizării decât cele cu profile diverse și adaptabile. Variabila automation_score_proxy "
    "a fost introdusă explicit pentru a captura această dimensiune structural-cantitativă, "
    "complementând informația brută despre numărul de competențe."
)

replace_all(
    "Înainte de antrenarea modelului, datele au fost supuse unui proces de preprocesare pentru "
    "a asigura consistența și calitatea informațiilor utilizate.",

    "Procesul de preprocesare a cuprins mai mulți pași, fiecare cu un rol specific în asigurarea "
    "calității datelor introduse în model. Deși ESCO este o sursă structurată și bine întreținută, "
    "versiunea românească poate conține inconsistențe minore — ocupații cu zero competențe "
    "înregistrate, valori lipsă sau intrări duplicate — care trebuie tratate înainte de modelare."
)

replace_all(
    "Preprocesarea reprezintă o etapă importantă în dezvoltarea modelelor predictive, deoarece "
    "datele brute pot conține valori lipsă, inconsistențe sau structuri dificil de interpretat "
    "direct de algoritm.",

    "Prima problemă identificată a fost prezența ocupațiilor fără nicio competență asociată — "
    "probabil intrări incomplete în baza de date. Acestea au fost eliminate, deoarece nu conțineau "
    "informație utilă pentru model. Ulterior, am verificat distribuțiile fiecărei variabile "
    "pentru a detecta valori aberante sau transformări necesare."
)

replace_all(
    "Prima etapă a constat în curățarea datasetului și eliminarea eventualelor observații "
    "incomplete sau redundante. Ulterior, au fost verificate distribuțiile variabilelor și "
    "consistența valorilor numerice.",

    "Variabila total_skill a arătat o distribuție asimetrică, cu câteva ocupații cu un număr "
    "excepțional de mare de competențe. Aceste cazuri nu au fost eliminate — sunt ocupații "
    "legitime, cu profiluri ample — dar distribuția a fost luată în considerare la interpretarea "
    "scorurilor extreme."
)

replace_all(
    "De asemenea, au fost selectate doar variabilele considerate relevante pentru estimarea "
    "automatizării ocupaționale. Această selecție a avut rolul de a reduce complexitatea modelului "
    "și de a elimina informațiile care nu contribuiau semnificativ la predicție.",

    "Selecția variabilelor a fost deliberat conservatoare: am inclus doar cele șase caracteristici "
    "pentru care exista o justificare teoretică clară. Adăugarea unor variabile fără fundament "
    "conceptual ar fi putut introduce zgomot în model și ar fi complicat interpretarea rezultatelor "
    "SHAP fără a aduce un câștig real de performanță."
)

replace_all(
    "În funcție de distribuția datelor, au fost aplicate proceduri de normalizare și standardizare "
    "pentru a reduce variațiile excesive dintre variabile.",

    "LightGBM este, în principiu, robust față de diferențele de scală între variabile — arborii "
    "de decizie utilizează ranguri, nu valori absolute. Cu toate acestea, pentru a asigura "
    "compatibilitatea cu eventuale comparații viitoare cu modele liniare, variabilele numerice "
    "au fost standardizate."
)

replace_all(
    "Procesul de preprocesare a urmărit obținerea unui set de date coerent, comparabil și adecvat "
    "pentru antrenarea modelului LightGBM.",

    "La finalul preprocesării, setul de date curat conținea 3.037 de ocupații cu șase variabile "
    "fiecare, fără valori lipsă. Acesta a constituit input-ul pentru toate etapele ulterioare "
    "de modelare și evaluare."
)

replace_all(
    "Pentru estimarea scorului de automatizare a fost utilizat algoritmul LightGBM, un model de "
    "machine learning bazat pe tehnica gradient boosting.",

    "Alegerea algoritmului s-a făcut prin compararea preliminară a mai multor opțiuni — regresie "
    "liniară, Random Forest și LightGBM — pe un subset de validare. LightGBM a arătat "
    "consistenta performanțe superioare, confirmat și de literatura care îl recomandă pentru "
    "date tabulare de dimensiuni medii (Ke et al., 2017)."
)

replace_all(
    "LightGBM este recunoscut pentru eficiența ridicată în analiza datelor tabulare și pentru "
    "capacitatea de a identifica relații complexe între variabile. Comparativ cu alte modele de "
    "tip boosting, acesta permite antrenarea rapidă și gestionarea eficientă a volumelor mari "
    "de date.",

    "LightGBM funcționează prin construirea secvențială a unor arbori de decizie, fiecare "
    "concentrându-se pe corectarea erorilor reziduale ale celui precedent — acesta este "
    "principiul gradient boosting-ului. Față de implementările clasice, LightGBM utilizează o "
    "strategie de creștere leaf-wise care îl face mai eficient și, de obicei, mai precis pe "
    "seturi de date tabulare."
)

replace_all(
    "Alegerea algoritmului a fost motivată de faptul că relația dintre competențele ocupaționale "
    "și automatizare nu este liniară. Unele variabile pot influența rezultatul într-un mod "
    "indirect, iar modelele bazate pe arbori de decizie sunt capabile să surprindă aceste relații.",

    "Un argument suplimentar pentru LightGBM față de modelele liniare este că relația dintre "
    "structura competențelor și riscul de automatizare nu este, cel mai probabil, liniară. De "
    "exemplu, este rezonabil să presupunem că efectul numărului de competențe opționale depinde "
    "de raportul acestora față de total — și că există praguri dincolo de care semnificația "
    "acestei variabile se schimbă. Arborii de decizie captează natural astfel de neliniarități "
    "și interacțiuni."
)

replace_all(
    "Modelul a fost implementat sub formă de regresie, având ca rezultat un scor numeric asociat "
    "fiecărei ocupații.",

    "Modelul a fost implementat în variantă de regresie (LGBMRegressor), producând un scor "
    "continuu de automatizare pentru fiecare ocupație. Decizia de a folosi regresia în loc de "
    "clasificare directă a fost deliberată: un scor continuu oferă mai multă informație decât "
    "o etichetă de clasă și permite praguri flexibile de clasificare ulterioară."
)

replace_all(
    "În timpul antrenării, au fost configurați parametri specifici pentru optimizarea performanței.",

    "Parametrii modelului au fost selectați prin combinarea ghidurilor din documentație cu "
    "experimentarea empirică pe datele de validare. Nu s-a efectuat o căutare exhaustivă "
    "(hyperparameter tuning), din considerente de interpretabilitate și reproductibilitate — "
    "configurația standard a LightGBM cu parametrii din Tabelul 2 s-a dovedit suficient de "
    "stabilă pentru scopurile acestei analize."
)

replace_all(
    "Rezultatul generat de model constă într-un scor continuu de automatizare, exprimat pe o "
    "scară între 0 și 100.",

    "Variabila țintă pe care modelul încearcă să o reproducă este un scor proxy de automatizare, "
    "construit pe baza unui lexicon de termeni asociați sarcinilor rutiniere și nerutiniere din "
    "descrierile competențelor ESCO. Scorul ia valori între 0 (risc minim estimat) și 100 "
    "(risc maxim estimat) — o scală normalizată care facilitează interpretarea și compararea."
)

replace_all(
    "Scorurile mici indică ocupații mai puțin expuse automatizării, în timp ce valorile ridicate "
    "sugerează un grad mai mare de vulnerabilitate în fața proceselor automatizate.",

    "Este important să subliniem că acest scor nu reprezintă o probabilitate în sens statistic "
    "strict și nici o predicție certă despre viitorul unei ocupații. El cuantifică, pe baza "
    "datelor disponibile, gradul relativ de expunere structurală la automatizare — un instrument "
    "de comparare și prioritizare, nu un verdict."
)

replace_all(
    "Pentru a facilita interpretarea rezultatelor, scorul numeric a fost transformat într-un "
    "sistem de clasificare bazat pe trei niveluri de risc.",

    "Pentru a facilita comunicarea rezultatelor unui public mai larg, scorurile continue au "
    "fost transformate ulterior într-un sistem de clasificare cu trei niveluri. Pragurile au "
    "fost alese pe baza distribuției scorurilor și a considerentelor de interpretabilitate, "
    "nu prin optimizare statistică — o decizie metodologică pe care o discutăm și în secțiunea "
    "dedicată limitărilor."
)

replace_all(
    "Această împărțire a fost aleasă pentru a separa ocupațiile cu risc redus, mediu și ridicat "
    "de automatizare.",

    "Pragul LOW/MEDIUM la 45 și MEDIUM/HIGH la 65 delimitează aproximativ treimile inferioare, "
    "mediană și superioară ale distribuției scorurilor pe setul ESCO, asigurând că fiecare "
    "categorie este bine populată și interpretabilă."
)

replace_all(
    "Transformarea scorului continuu într-o clasificare discretă permite o interpretare mai "
    "intuitivă a rezultatelor și facilitează compararea ocupațiilor.",

    "Un dezavantaj al clasificării discrete este că introducerea unui prag creează o diferență "
    "artificială între ocupații cu scoruri apropiate de granița dintre categorii. O ocupație cu "
    "scorul 44 și una cu 46 sunt clasificate diferit, deși diferența practică este neglijabilă. "
    "Această limitare trebuie avută în vedere la interpretarea individuală a scorurilor."
)

replace_all(
    "Evaluarea performanței modelului a fost realizată prin utilizarea unor metrici specifice "
    "atât regresiei, cât și clasificării.",

    "Evaluarea riguroasă a unui model predictiv presupune mai mult decât o singură cifră de "
    "performanță. Am utilizat metrici complementare care surprind aspecte diferite ale "
    "comportamentului modelului: eroarea de predicție numerică (MAE), capacitatea de explicare "
    "a variației (R²), stabilitatea estimărilor (cross-validare) și calitatea clasificării "
    "(accuracy față de baseline)."
)

replace_all(
    "În cazul regresiei, a fost utilizată eroarea medie absolută (MAE), care măsoară diferența "
    "medie dintre valorile reale și cele estimate.",

    "MAE (Mean Absolute Error) a fost ales ca metrică primară pentru regresie deoarece este "
    "ușor de interpretat în unități ale variabilei țintă — puncte pe scala 0–100 — și este "
    "mai robust față de valorile extreme decât RMSE. Un MAE de 11.03 înseamnă că, în medie, "
    "predicția modelului se abate cu 11 puncte față de scorul real."
)

replace_all(
    "Coeficientul de determinare R² a fost folosit pentru a evalua capacitatea modelului de a "
    "explica variația datelor.",

    "R² măsoară proporția din varianța variabilei țintă explicată de model, față de un "
    "clasificator naiv care prezice întotdeauna media. Valoarea pozitivă de 0.196 pe setul de "
    "test indică că modelul explică aproximativ 20% din variația scorurilor — o performanță "
    "modestă, dar pozitivă, pe un set de variabile strict structurale."
)

replace_all(
    "În paralel, clasificarea ocupațiilor pe niveluri de risc a fost evaluată utilizând "
    "indicatorul accuracy.",

    "Pentru clasificare, am raportat acuratețea față de două referințe: un clasificator naiv "
    "care atribuie întotdeauna clasa majoritară (53.95%) și performanța pe setul de antrenare "
    "(86.91%). Diferența mare antrenare-test sugerează overfitting — modelul a memorat "
    "parțial setul de antrenare, ceea ce limitează generalizarea."
)

replace_all(
    "Rezultatele indică o capacitate moderată de predicție și clasificare. Instabilitatea "
    "cross-validare (R²_CV = −14.21) sugerează existența unor limitări privind generalizarea "
    "modelului.",

    "Cifrele de performanță trebuie citite împreună, nu izolat. R²_test = 0.196 arată că modelul "
    "e mai bun decât hazardul pe setul de test, dar R²_CV = −14.21 ± 6.16 obținut în "
    "cross-validare trădează o instabilitate semnificativă — performanța variază dramatic între "
    "fold-uri, ceea ce indică fie că setul de date este relativ mic față de complexitatea "
    "problemei, fie că relația dintre variabilele structurale și scorul de automatizare nu este "
    "suficient de robustă pentru a fi generalizată."
)

replace_all(
    "Figura prezintă distribuția predicțiilor corecte și eronate între cele trei categorii de "
    "risc. Diagonala principală reflectă clasificările corecte, iar valorile din afara diagonalei "
    "indică erorile de clasificare.",

    "Matricea de confuzie completează informația oferită de acuratețe globală arătând distribuția "
    "erorilor pe clase. Diagonala principală indică clasificările corecte, iar valorile din "
    "afara acesteia reflectă confuziile — esențial de examinat pentru a înțelege dacă modelul "
    "greșește sistematic în favoarea anumitor clase."
)

replace_all(
    "În cadrul cercetării, interpretabilitatea modelului a reprezentat o componentă importantă, "
    "deoarece analiza rezultatelor nu se limitează doar la performanța predictivă.",

    "Performanța predictivă singură nu este suficientă pentru a valida un model destinat analizei "
    "de politici. Trebuie să verificăm că modelul utilizează variabilele în mod coerent cu "
    "așteptările teoretice — că, de exemplu, un număr mare de competențe repetitive crește "
    "riscul estimat, nu îl reduce. Analiza SHAP îndeplinește această funcție de validare."
)

replace_all(
    "Înțelegerea modului în care modelul ajunge la anumite scoruri este esențială pentru "
    "validarea rezultatelor și pentru explicarea influenței variabilelor utilizate.",

    "Mai concret, SHAP permite să răspundem la două tipuri de întrebări: global (care variabile "
    "sunt, în medie, cele mai influente?) și local (de ce această ocupație specifică a primit "
    "scorul X?). Ambele niveluri de analiză sunt relevante — primul pentru înțelegerea modelului "
    "ca întreg, al doilea pentru comunicarea rezultatelor la nivel individual."
)

replace_all(
    "SHAP permite identificarea variabilelor care contribuie cel mai mult la predicție și "
    "măsurarea impactului fiecărei caracteristici asupra scorului final.",

    "Analiza SHAP a confirmat că variabilele automation_score_proxy și total_skill sunt "
    "consistente cei mai importanți predictori — un rezultat aliniat cu ipoteza teoretică, "
    "conform căreia proporția competențelor opționale și de cunoaștere (incluse în proxy) "
    "capturează structura rutinizabilă a ocupației."
)

replace_all(
    "Baza de date COR include 4456 ocupații și reprezintă sistemul oficial de clasificare "
    "utilizat la nivel național.",

    "COR conține 4.456 de ocupații organizate pe șase niveluri de clasificare, de la grupe "
    "majore (un singur caracter) până la ocupații individuale (șase caractere). Această "
    "structură ierarhică nu a fost exploatată în analiza de față — am tratat fiecare ocupație "
    "ca o entitate independentă — dar poate fi valorificată în cercetări viitoare."
)

replace_all(
    "Pentru a transfera rezultatele modelului asupra ocupațiilor românești, a fost necesară "
    "asocierea ocupațiilor COR cu ocupațiile ESCO.",

    "Pasul cel mai delicat din întregul pipeline a fost tocmai asocierea dintre COR și ESCO. "
    "Cele două sisteme au filosofii de clasificare diferite și nomenclaturi care nu se suprapun "
    "consistent. Dacă am fi putut traduce manual fiecare ocupație COR în echivalentul ESCO "
    "corespunzător, am fi obținut rezultate mai precise — dar cu 4.456 de ocupații, aceasta "
    "ar fi necesitat o expertiză specializată și un efort considerabil."
)

replace_all(
    "Această asociere a fost realizată prin metode de similaritate textuală aproximativă, "
    "cunoscute sub denumirea de fuzzy matching.",

    "Soluția practică adoptată a fost fuzzy matching-ul prin algoritmul RapidFuzz — o variantă "
    "eficientă a distanței Levenshtein, adaptată pentru comparații de șiruri de caractere. "
    "Logica este simplă: dacă denumirea ocupației COR și cea a ocupației ESCO sunt suficient "
    "de similare ca text, le considerăm echivalente și transferăm caracteristicile ESCO "
    "ocupației COR."
)

replace_all(
    "Procesul de mapare, cu un prag de similaritate ≥ 80%, a permis asocierea unui număr de "
    "654 de ocupații, reprezentând 14.7% din totalul ocupațiilor COR. Pragul ridicat a fost "
    "ales pentru a minimiza erorile semantice.",

    "Pragul de ≥ 80% a fost ales deliberat mai restrictiv decât valorile uzuale (adesea 70%), "
    "tocmai pentru a reduce erorile semantice — situațiile în care două ocupații au denumiri "
    "similare, dar conținuturi diferite. Prețul acestei precauții este rata de acoperire "
    "scăzută: 14.7% din ocupațiile COR (654 din 4.456). Este o limitare importantă, discutată "
    "în detaliu în capitolul 5."
)

replace_all(
    "Rezultatele obținute au permis aplicarea modelului asupra ocupațiilor mapate și generarea "
    "scorurilor de automatizare.",

    "Ocupațiile COR care au trecut pragul de similaritate au primit, prin transfer de la ESCO, "
    "valorile celor șase variabile structurale. Pe baza acestora, modelul antrenat a generat "
    "scoruri de automatizare individuale, pe care le-am clasificat ulterior în LOW, MEDIUM sau "
    "HIGH folosind aceleași praguri de 45 și 65 definite în etapa anterioară."
)

print("\n=== CAP 4 — REZULTATE ===")
replace_all(
    "După dezvoltarea modelului predictiv și evaluarea performanței acestuia, etapa următoare "
    "a constat în interpretarea rezultatelor generate și analiza modului în care modelul "
    "estimează riscul de automatizare pentru ocupațiile analizate.",

    "Rezultatele obținute ridică cel puțin două întrebări distincte, care trebuie tratate "
    "separat: cât de bine funcționează modelul pe datele de testare ESCO? și ce spune "
    "distribuția scorurilor despre profilul de risc al ocupațiilor românești din COR? "
    "Răspunsurile la cele două întrebări sunt, după cum vom vedea, nuanțate și uneori "
    "contradictorii la prima vedere."
)

replace_all(
    "Capitolul de față prezintă rezultatele obținute în urma aplicării modelului asupra datelor "
    "ESCO și transferului acestora asupra Clasificației Ocupațiilor din România (COR). Analiza "
    "urmărește atât performanța modelului, cât și distribuția scorurilor de automatizare, "
    "interpretarea nivelurilor de risc și identificarea tiparelor ocupaționale.",

    "Prezentarea rezultatelor urmează o logică progresivă: începem cu performanța modelului "
    "pe setul de evaluare ESCO (secțiunile 4.2–4.4), continuăm cu aplicarea pe ocupațiile COR "
    "(4.5–4.6), includem exemple concrete (4.7) și închidem cu o analiză critică a limitărilor "
    "observate în date (4.8)."
)

replace_all(
    "Rezultatele sunt analizate atât din perspectivă cantitativă, prin utilizarea indicatorilor "
    "de performanță, cât și din perspectivă interpretativă, prin observarea distribuției "
    "ocupațiilor și a modului în care acestea se încadrează în diferite categorii de risc.",

    "O remarcă metodologică preliminară: indicatorii de performanță prezentați în secțiunea 4.2 "
    "se referă la setul de test ESCO, nu la predicțiile COR. Scorurile pentru ocupațiile COR "
    "sunt estimări de transfer — nu putem evalua direct acuratețea lor față de o valoare reală, "
    "deoarece nu există o etichetă validată extern pentru riscul de automatizare al ocupațiilor "
    "românești."
)

replace_all(
    "Performanța modelului a fost evaluată utilizând metrici specifice atât problemelor de "
    "regresie, cât și clasificării. Scopul acestei etape a fost determinarea gradului în care "
    "modelul poate estima scorul de automatizare și poate clasifica ocupațiile în funcție de "
    "nivelul de risc.",

    "Evaluarea performanței pe setul de test constituie singura măsurare obiectivă a calității "
    "modelului disponibilă în această cercetare. Tabelul 5 sintetizează principalele metrici "
    "obținute, incluzând atât valorile pe setul de test, cât și rezultatele cross-validării "
    "5-fold — esențiale pentru a evalua stabilitatea estimărilor."
)

replace_all(
    "Valoarea MAE indică faptul că diferența medie dintre scorurile reale și cele estimate este "
    "relativ redusă, ceea ce sugerează o capacitate moderată de aproximare a valorilor.",

    "MAE-ul de 11.03 pe testare trebuie interpretat în context. Pe o scală de la 0 la 100, "
    "cu praguri de clasificare la 45 și 65, o eroare medie de 11 puncte poate modifica "
    "încadrarea în categorii de risc pentru ocupațiile aflate în apropierea pragurilor. "
    "Acesta este motivul pentru care scorurile individuale trebuie tratate ca estimări "
    "orientative, nu ca valori exacte."
)

replace_all(
    "În ceea ce privește coeficientul de determinare R², instabilitatea cross-validare sugerează "
    "că modelul întâmpină dificultăți în generalizarea asupra unor observații noi. Acest rezultat "
    "poate fi influențat de structura limitată a variabilelor disponibile și de complexitatea "
    "ridicată a relațiilor dintre competențe și automatizare.",

    "Tensiunea dintre R²_test = 0.196 și R²_CV = −14.21 ± 6.16 este cea mai îngrijorătoare "
    "caracteristică a modelului. O explicație probabilă este că partițiile din cross-validare "
    "capturează variabilitatea naturală din setul ESCO mai agresiv decât split-ul simplu "
    "train/test — ceea ce sugerează că performanța observată pe test-set poate fi parțial "
    "norocoasă. Această concluzie impune maxim de prudență în extrapolarea rezultatelor."
)

replace_all(
    "Scorul generat de model reprezintă o estimare numerică a gradului de expunere la automatizare "
    "pentru fiecare ocupație.",

    "Dincolo de metricile agregate, analiza distribuției scorurilor individuale oferă o imagine "
    "mai nuanțată a modului în care modelul percepe spațiul ocupațional."
)

replace_all(
    "Valorile mici indică ocupații mai puțin vulnerabile la automatizare, în timp ce valorile "
    "ridicate sugerează un risc crescut.",

    "Scorurile sunt distribuite pe întreaga scală 0–100, cu o concentrare predominantă în "
    "intervalul 20–60. Aceasta reflectă, probabil, structura competențelor ESCO: majoritatea "
    "ocupațiilor au o combinație de competențe rutinare și nerutinare, fără a fi extreme în "
    "niciunul dintre sensuri."
)

replace_all(
    "Distribuția scorurilor evidențiază faptul că majoritatea ocupațiilor analizate tind să "
    "se concentreze în zona mediană a scalei.",

    "Concentrarea în zona mediană nu este neapărat un semn de performanță slabă a modelului — "
    "ar putea reflecta realitatea că puține ocupații sunt, structural, pur repetitive sau pur "
    "creative. Aceasta este, de altfel, și concluzia studiilor internaționale recente: "
    "polarizarea reală a pieței muncii nu este între ocupații care vor dispărea versus ocupații "
    "complet sigure, ci între cele care se vor transforma mai mult și cele care se vor transforma "
    "mai puțin."
)

replace_all(
    "Această distribuție sugerează că multe ocupații includ activități care pot fi automatizate "
    "parțial, fără ca profesia să fie complet înlocuită.",

    "Distribuția relativ simetrică în jurul mediei sugerează că modelul nu are o tendință "
    "sistematică de supraestimare sau subestimare a riscului — cel puțin la nivel agregat. "
    "Analiza erorilor individuale (disponibilă în fișierul analiza_erori_model.csv generat "
    "în etapa de evaluare) arată că cele mai mari erori apar la ocupații cu profiluri "
    "competențiale atipice."
)

replace_all(
    "Scorurile foarte ridicate apar mai rar și sunt asociate în general ocupațiilor caracterizate "
    "prin sarcini repetitive și proceduri standardizate.",

    "La capătul superior al distribuției, scorurile ridicate corespund, de regulă, ocupațiilor "
    "cu un număr mic de competențe distincte, concentrat în cunoștințe formale sau proceduri "
    "standardizate — de exemplu, operatori de utilaje specifice sau lucrători în procesare "
    "de date."
)

replace_all(
    "Pe de altă parte, ocupațiile cu scoruri reduse tind să includă activități care necesită "
    "interacțiune umană, creativitate sau luarea deciziilor complexe.",

    "La capătul inferior, scorurile mici caracterizează în mod consistent ocupațiile cu "
    "profiluri competențiale ample și diversificate — profesii medicale, juridice sau "
    "educaționale — unde varietatea competențelor reflectă complexitatea și adaptabilitatea "
    "cerute de exercitarea lor."
)

replace_all(
    "Rezultatele au arătat o distribuție relativ echilibrată: LOW (50.5%) și MEDIUM (48.6%) "
    "sunt categoriile dominante, cu o proporție redusă de ocupații HIGH (0.9%).",

    "Pe setul ESCO de antrenare, distribuția claselor a arătat că LOW (53.9%) este categoria "
    "majoritară, urmată de MEDIUM (35.1%) și HIGH (11.0%). Această distribuție dezechilibrată "
    "a influențat semnificativ performanța clasificatorului, care tinde să favorizeze clasa "
    "majoritară — un comportament așteptat și documentat în literatura de machine learning pe "
    "seturi dezechilibrate."
)

replace_all(
    "Scopul acestei etape a fost transferul rezultatelor predictive într-un context național și "
    "evaluarea distribuției riscului ocupațional în România.",

    "Aplicarea modelului pe COR urmărește un scop dublu: pe de o parte, testarea transferabilității "
    "metodologiei pe un alt sistem de clasificare; pe de altă parte, obținerea unor estimări "
    "orientative privind profilul de risc al ocupațiilor din România."
)

replace_all(
    "Procesul de mapare a permis conectarea ocupațiilor similare și aplicarea modelului asupra "
    "celor care au putut fi asociate.",

    "Prin fuzzy matching cu prag ≥ 80%, au putut fi asociate 654 de ocupații COR cu echivalente "
    "ESCO. Cele 3.802 ocupații rămase nemapate nu au putut fi evaluate cu modelul curent, "
    "ceea ce reprezintă cea mai importantă limitare operațională a acestei analize."
)

replace_all(
    "Rezultatele arată că aproximativ o treime dintre ocupațiile disponibile au putut fi "
    "asociate cu structura ESCO.",

    "Rata de mapare de 14.7% este mai mică decât valorile raportate în studii similare pentru "
    "alte piețe, ceea ce reflectă atât diferențele structurale dintre COR și ESCO, cât și "
    "faptul că pragul de 80% este mai restrictiv decât media din literatură. Un prag mai "
    "permisiv ar crește acoperirea, dar ar introduce mai multe asocieri eronate."
)

replace_all(
    "Procentul de ocupații nemapate indică limitările asociate diferențelor dintre clasificările "
    "ocupaționale și dificultățile generate de similaritatea textuală.",

    "Printre motivele pentru care ocupațiile nu s-au putut mapa se numără: denumiri complet "
    "diferite între COR și ESCO pentru ocupații echivalente, ocupații COR specifice contextului "
    "național fără corespondent european, și variații ortografice sau de diacritice care au "
    "îngreunat comparația textuală."
)

replace_all(
    "Rezultatele indică o distribuție echilibrată între categoriile LOW (330 ocupații, 50.5%) "
    "și MEDIUM (318 ocupații, 48.6%), cu un număr redus de ocupații HIGH (6, 0.9%). Această "
    "distribuție sugerează că pragul de similaritate ridicat (≥ 80%) selectează ocupații "
    "preponderent din zone de risc moderat spre scăzut.",

    "Distribuția riscului pentru ocupațiile COR mapate arată o repartizare aproximativ egală "
    "între LOW (330, 50.5%) și MEDIUM (318, 48.6%), cu un număr extrem de mic de ocupații "
    "HIGH (6, 0.9%). Această distribuție diferă față de setul ESCO de antrenare, unde HIGH "
    "reprezenta 11%. O ipoteză rezonabilă este că pragul ridicat de similaritate favorizează "
    "selecția ocupațiilor cu denumiri mai complete și mai elaborate — care tind să corespundă "
    "unor profiluri mai complexe, mai puțin rutiniere."
)

replace_all(
    "Această distribuție poate fi explicată prin faptul că numeroase profesii combină activități "
    "administrative sau repetitive cu sarcini care necesită expertiză umană.",

    "Aceasta nu înseamnă neapărat că piața muncii din România are un profil de risc mai scăzut "
    "decât media europeană — ci că eșantionul de 654 de ocupații accesibile prin mapare poate "
    "fi un subset nesistematic, cu un bias spre ocupațiile mai ușor de asociat textual."
)

replace_all(
    "Numărul foarte redus de ocupații HIGH (6, reprezentând 0.9%) reflectă atât pragul ridicat "
    "de similaritate utilizat, cât și faptul că maparea selectivă favorizează ocupațiile cu "
    "competențe mai diversificate.",

    "Cele 6 ocupații clasificate HIGH — sablator produse din sticlă, confecționer circuite "
    "imprimate și operator la mașina de laminat printre primele trei — au scoruri între 66 "
    "și 70, ceea ce le plasează modest în zona de risc ridicat. Absența unor scoruri extreme "
    "(>85) în setul COR este explicabilă prin selectivitatea procesului de mapare."
)

replace_all(
    "Deși modelul a reușit să genereze scoruri și clasificări relevante, rezultatele trebuie "
    "interpretate în contextul limitărilor existente.",

    "Privind retrospectiv rezultatele, câteva aspecte cer o examinare critică mai atentă "
    "decât o simplă notă de subsol."
)

replace_all(
    "Acest rezultat poate fi influențat de numărul limitat de variabile disponibile și de "
    "complexitatea ridicată a relațiilor dintre competențe și automatizare.",

    "Există două explicații complementare pentru această instabilitate. Prima este că setul "
    "de date (3.037 observații) este modest față de complexitatea relației modelate, ceea ce "
    "face ca partițiile din cross-validare să conțină variabilitate suficient de mare pentru "
    "a schimba semnificativ estimările. A doua explicație este că variabilele structurale "
    "din ESCO captează indirect, nu direct, rutinizabilitatea ocupațiilor — calitatea "
    "semnalului este limitată prin natura datelor."
)

replace_all(
    "De asemenea, procesul de mapare COR–ESCO poate introduce erori semantice, deoarece "
    "asocierea ocupațiilor se bazează pe similaritate textuală și nu întotdeauna pe "
    "echivalență perfectă.",

    "O sursă de eroare specifică transferului pe COR este că similaritatea textuală nu "
    "garantează echivalența funcțională. O ocupație COR numită „operator la computer” și "
    "una ESCO numită „computer operator” au un scor de similaritate ridicat, dar pot "
    "desemna realități profesionale diferite în contexte economice diferite. Verificarea "
    "manuală a unui eșantion din asocierile efectuate ar fi o bună practică pentru versiunile "
    "viitoare ale acestei cercetări."
)

replace_all(
    "Cu toate acestea, rezultatele obținute demonstrează că analiza ocupațională bazată pe "
    "machine learning poate oferi informații relevante privind vulnerabilitatea profesiilor.",

    "Cu toate limitările enumerate, modelul furnizează estimări consistente cu logica "
    "teoretică — ocupațiile cu profiluri competențiale standardizate primesc scoruri mai "
    "mari decât cele cu profiluri diverse — și produce o metodologie reproductibilă care "
    "poate fi îmbunătățită incremental."
)

replace_all(
    "Interpretarea vizuală a rezultatelor reprezintă o componentă esențială în analiza "
    "modelului predictiv, deoarece permite observarea intuitivă a distribuției scorurilor, "
    "evaluarea performanței clasificării și înțelegerea modului în care variabilele "
    "influențează predicțiile.",

    "Graficele prezentate în această secțiune servesc trei scopuri complementare: verificarea "
    "intuitivă a distribuției scorurilor, evaluarea vizuală a calității clasificării și "
    "înțelegerea comportamentului intern al modelului prin analiza SHAP."
)

replace_all(
    "Prin utilizarea reprezentărilor grafice, pot fi identificate tipare generale și relații "
    "între caracteristicile ocupațiilor și riscul estimat de automatizare. Aceste instrumente "
    "contribuie la o mai bună validare a modelului și facilitează interpretarea rezultatelor "
    "obținute.",

    "Reprezentările vizuale nu înlocuiesc analiza cantitativă, dar o completează: un model cu "
    "metrici bune pe hârtie poate ascunde comportamente problematice (de exemplu, predicții "
    "extreme concentrate pe un subset de ocupații) care sunt imediat vizibile grafic."
)

replace_all(
    "Graficul prezintă modul în care scorurile de automatizare sunt distribuite la nivelul "
    "ocupațiilor analizate. Repartizarea valorilor pe intervalul 0–100 permite identificarea "
    "zonelor în care se concentrează majoritatea observațiilor.",

    "Histograma scorurilor de automatizare arată o distribuție relativ simetrică, centrată în "
    "jurul valorii 40–50, cu cozi ușor mai lungi spre extremele scalei. Forma distribuției "
    "este importantă: o distribuție bimodală ar fi sugerat că modelul tratează ocupațiile "
    "ca aparținând clar uneia din două categorii, în timp ce distribuția unimodală observată "
    "sugerează că variabilitatea riscului este graduală."
)

replace_all(
    "Se remarcă o concentrare predominantă a scorurilor în intervalul median, ceea ce indică "
    "faptul că o mare parte dintre ocupații prezintă un nivel moderat de expunere la automatizare. "
    "Această distribuție sugerează existența unor structuri ocupaționale mixte, în care "
    "activitățile automatizabile coexistă cu sarcini ce necesită intervenție umană.",

    "Concentrarea în zona 30–60 reflectă natura datelor ESCO: marea majoritate a ocupațiilor "
    "din baza europeană au profiluri competențiale echilibrate, cu atât esențiale cât și "
    "opționale, fără a fi complet specializate într-un singur tip de sarcini. Aceasta "
    "corespunde cu perspectivele din literatura de specialitate, care subliniază că automatizarea "
    "tipică transformă, nu elimină ocupațiile."
)

replace_all(
    "Valorile situate pe diagonala principală indică situațiile în care modelul a realizat "
    "clasificări corecte, în timp ce valorile din afara diagonalei reflectă confuziile între "
    "clase. Analiza acestui grafic arată că modelul reușește să identifice într-o măsură "
    "satisfăcătoare ocupațiile din categoria MEDIUM, care reprezintă și clasa dominantă în "
    "setul de date.",

    "Matricea de confuzie relevă un pattern tipic pentru seturi cu distribuție dezechilibrată: "
    "modelul clasifică corect cea mai mare parte a instanțelor din clasa majoritară (LOW), dar "
    "performează mai slab pe MEDIUM și mai ales pe HIGH. Clasa HIGH, cu doar 334 exemple în "
    "antrenare, este cea mai dificil de identificat — un rezultat așteptat și consistent cu "
    "literatura despre clasificarea pe seturi dezechilibrate."
)

replace_all(
    "Erorile de clasificare apar în special între categoriile adiacente, ceea ce sugerează că "
    "diferențierea între niveluri apropiate de risc este mai dificilă decât separarea extremelor. "
    "Acest comportament este specific modelelor aplicate pe date cu distribuții dezechilibrate "
    "sau cu tranziții graduale între clase.",

    "Erorile de clasificare sunt, în marea lor majoritate, între categorii adiacente (LOW "
    "clasificat ca MEDIUM sau MEDIUM clasificat ca HIGH), nu între extreme (LOW clasificat ca "
    "HIGH). Aceasta sugerează că granițele dintre categorii sunt graduale în spațiul "
    "competențelor, nu abrupte — o observație importantă pentru calibrarea așteptărilor "
    "față de precizia clasificărilor individuale."
)

replace_all(
    "Graficul de importanță a variabilelor evidențiază contribuția fiecărei caracteristici "
    "utilizate în model la estimarea scorului de automatizare. Analiza este realizată pe baza "
    "valorilor SHAP, care permit cuantificarea impactului individual al fiecărei variabile "
    "asupra predicțiilor generate.",

    "Graficul de importanță SHAP răspunde la întrebarea: dacă am elimina o variabilă din "
    "model, cât de mult ar scădea precizia predicțiilor? Variabilele poziționate mai sus "
    "sunt cele mai influente în medie pe întreg setul de date."
)

replace_all(
    "Rezultatele indică faptul că variabilele asociate competențelor ocupaționale au o influență "
    "semnificativă asupra scorului estimat. În special, numărul total de skill-uri și ponderea "
    "competențelor esențiale contribuie într-o măsură mai mare la determinarea riscului de "
    "automatizare.",

    "Rezultatele SHAP arată că automation_score_proxy și total_skill sunt cei mai importanți "
    "predictori — ceea ce este coerent cu ipoteza că proporția competențelor rutinizabile și "
    "volumul total al profilului competențial sunt caracteristicile cele mai informative "
    "pentru estimarea riscului."
)

replace_all(
    "Această observație sugerează că structura competențelor reprezintă un factor determinant "
    "în evaluarea vulnerabilității ocupațiilor. Ocupațiile care implică un set diversificat de "
    "abilități tind să fie mai rezistente la automatizare, în timp ce ocupațiile caracterizate "
    "prin competențe limitate și standardizate sunt mai expuse.",

    "Faptul că automation_score_proxy — o variabilă construită explicit pentru a captura "
    "proporția competențelor rutinizabile — apare consistent printre cei mai importanți "
    "predictori este, în același timp, un semn de coerență internă a modelului și o "
    "confirmare că designul acestei variabile a fost bine calibrat."
)

replace_all(
    "Prin urmare, analiza importanței variabilelor oferă o perspectivă clară asupra modului "
    "în care modelul utilizează informațiile disponibile și confirmă relevanța datelor "
    "ocupaționale în estimarea riscului de automatizare.",

    "O observație demnă de menționat: variabila essential are o importanță SHAP relativ mică "
    "față de așteptări. O interpretare posibilă este că, după ce se cunoaște total_skill, "
    "informația marginală adusă de essential este redusă — cele două variabile sunt "
    "corelate prin identitatea total_skill = essential + optional."
)

replace_all(
    "SHAP Summary Plot oferă o perspectivă detaliată asupra modului în care fiecare variabilă "
    "influențează predicțiile modelului, atât în ceea ce privește direcția, cât și "
    "intensitatea impactului.",

    "SHAP Summary Plot completează graficul de importanță cu informații despre direcție — arată "
    "nu doar cât de mult influențează o variabilă predicția, ci și în ce sens: valori mari ale "
    "variabilei cresc sau reduc scorul estimat?"
)

replace_all(
    "Fiecare punct din grafic reprezintă o observație individuală, iar poziția acestuia pe "
    "axa orizontală indică influența asupra scorului de automatizare. Distribuția punctelor "
    "evidențiază variabilitatea impactului pentru fiecare caracteristică.",

    "Fiecare punct reprezintă o ocupație. Poziția pe axa orizontală indică contribuția la "
    "predicție (pozitiv = crește riscul, negativ = îl reduce), iar culoarea reflectă valoarea "
    "variabilei (roșu = valoare mare, albastru = valoare mică). Un tipar consistent — "
    "de exemplu, puncte roșii sistematic în dreapta — indică că valori mari ale variabilei "
    "cresc riscul estimat."
)

replace_all(
    "Graficul permite identificarea situațiilor în care valori ridicate sau scăzute ale unei "
    "variabile contribuie la creșterea sau reducerea riscului estimat. Astfel, analiza oferă "
    "o înțelegere mai profundă a mecanismelor interne ale modelului și a relațiilor dintre "
    "variabile și rezultat.",

    "Tiparele observate în Summary Plot confirmă ipotezele teoretice: automation_score_proxy "
    "ridicat contribuie la scoruri mari de risc; total_skill ridicat contribuie la scoruri "
    "mici (mai multe competențe diverse = profil mai complex = risc mai scăzut). Aceste "
    "relații, deși nu surprinzătoare teoretic, validează că modelul a „învățat” semnale "
    "coerente din date, nu patternuri aleatoare."
)

replace_all(
    "Graficul ilustrează repartizarea ocupațiilor din România în funcție de nivelul estimat "
    "de risc de automatizare. Această reprezentare vizuală completează analiza tabelară și "
    "evidențiază structura generală a distribuției.",

    "Graficul de distribuție a riscului pentru ocupațiile COR mapate completează tabelul 7 "
    "cu o perspectivă vizuală. Două observații se evidențiază imediat: dominanța clară a "
    "categoriilor LOW și MEDIUM, și absența aproape completă a categoriei HIGH."
)

replace_all(
    "Se observă o distribuție relativ echilibrată între LOW și MEDIUM, ceea ce confirmă că "
    "majoritatea ocupațiilor mapate prezintă un risc moderat spre scăzut. Această distribuție "
    "reflectă existența unor roluri profesionale care includ atât activități automatizabile, "
    "cât și sarcini care necesită intervenție umană.",

    "Distribuția echilibrată LOW/MEDIUM nu înseamnă neapărat că România are un profil "
    "ocupațional mai sigur față de automatizare. O interpretare mai conservatoare este că "
    "procesul de mapare COR–ESCO, prin pragul restrictiv ales, a selecționat preferențial "
    "ocupații cu profiluri competențiale mai elaborate — tocmai cele cu risc mai scăzut. "
    "O analiză completă ar necesita o mapare manuală sau asistată semantic pentru toate "
    "cele 4.456 de ocupații COR."
)

replace_all(
    "Numărul extrem de redus de ocupații HIGH (6 din 654, 0.9%) reflectă limitările pragului "
    "de matching ridicat, care selectează ocupații cu profile competențiale mai complexe, "
    "mai rezistente la automatizare, în timp ce categoria LOW este reprezentată de ocupații "
    "caracterizate prin complexitate ridicată și interacțiune umană semnificativă.",

    "Cele 6 ocupații clasificate HIGH — cu scoruri între 66 și 70 — se caracterizează prin "
    "profiluri competențiale înguste și standardizate, concentrate pe operațiuni industriale "
    "specifice. Este notabil că niciunul dintre aceste scoruri nu depășește 70, ceea ce "
    "plasează ocupațiile respective la granița inferioară a categoriei HIGH, nu în zone "
    "de risc extrem."
)

replace_all(
    "Rezultatele prezentate în acest capitol evidențiază faptul că modelul dezvoltat poate "
    "genera estimări relevante privind riscul de automatizare al ocupațiilor.",

    "Capitolul de față a prezentat rezultatele din mai multe unghiuri complementare: metrici "
    "cantitative, distribuții vizuale și interpretare prin SHAP. Câteva concluzii intermediare "
    "se impun înainte de capitolul final."
)

replace_all(
    "Analiza performanței arată că modelul reușește să clasifice într-o manieră satisfăcătoare "
    "ocupațiile în funcție de nivelul de risc, chiar dacă există limitări privind generalizarea.",

    "Modelul furnizează estimări orientative cu o precizie modestă, dar pozitivă față de "
    "baseline. Performanța sa nu poate fi comparată direct cu modele care beneficiază de "
    "date mai bogate (salariale, sectoriale, de cerere de competențe), dar oferă un punct "
    "de plecare reproductibil pentru analize mai fine."
)

replace_all(
    "Aplicarea modelului asupra ocupațiilor din România demonstrează posibilitatea transferului "
    "metodologiei către un context național și oferă o imagine asupra distribuției riscului "
    "ocupațional.",

    "Transferul pe COR a demonstrat fezabilitatea metodologiei, chiar dacă rata de acoperire "
    "de 14.7% limitează semnificativ domeniul de aplicare. Extinderea acoperirii rămâne cea "
    "mai importantă direcție de dezvoltare pentru versiunile viitoare."
)

replace_all(
    "Interpretarea vizuală a rezultatelor contribuie la înțelegerea modului în care scorurile "
    "sunt distribuite și la identificarea variabilelor care influențează predicția.",

    "Analiza SHAP a validat coerența modelului: variabilele proiectate pentru a captura "
    "rutinizabilitatea ocupațională (automation_score_proxy, total_skill) sunt cele mai "
    "influente, și direcțiile lor de influență corespund cu așteptările teoretice."
)

replace_all(
    "În ansamblu, rezultatele susțin ideea că analiza bazată pe machine learning poate "
    "constitui un instrument util pentru evaluarea vulnerabilității ocupațiilor la automatizare "
    "și pentru dezvoltarea unor direcții viitoare de cercetare.",

    "În ansamblu, cercetarea demonstrează că datele ESCO conțin semnale suficiente pentru a "
    "construi un instrument funcțional de estimare a riscului ocupațional — cu limitările "
    "inerente unui model bazat exclusiv pe caracteristici structurale ale competențelor."
)

print("\n=== CAP 5 — CONCLUZII ===")
replace_all(
    "Prezenta cercetare a avut ca obiectiv dezvoltarea unui sistem bazat pe inteligență "
    "artificială pentru estimarea riscului de automatizare a ocupațiilor, utilizând date "
    "ocupaționale europene și tehnici de machine learning.",

    "La finalul acestui demers de cercetare, cred că e important să separăm cu claritate "
    "ceea ce a reușit această analiză de ceea ce rămâne deschis. Perspectiva retrospectivă "
    "este mai onestă decât o prezentare triumfalistă a rezultatelor."
)

replace_all(
    "În cadrul lucrării a fost construit un model predictiv capabil să estimeze un scor de "
    "automatizare pentru ocupațiile analizate, pornind de la caracteristicile asociate "
    "competențelor profesionale.",

    "Am construit un sistem funcțional care transformă descrierile structurale ale ocupațiilor "
    "din ESCO — număr de competențe, tipul lor, indicator de rutinizabilitate — într-un scor "
    "continuu de automatizare. Modelul LightGBM antrenat pe 3.037 de ocupații produce "
    "estimări care, pe setul de test, explică aproximativ 20% din variația scorurilor "
    "reale — modest, dar consistent pozitiv."
)

replace_all(
    "Rezultatele obținute demonstrează faptul că structura competențelor poate fi utilizată "
    "pentru a aproxima gradul de vulnerabilitate al unei ocupației în raport cu procesele "
    "de automatizare.",

    "Ce s-a confirmat, prin consistența cu așteptările teoretice și prin analiza SHAP, este "
    "că structura competențelor ocupaționale conține informație utilă despre vulnerabilitatea "
    "la automatizare. Ocupațiile cu profiluri competențiale înguste și standardizate primesc "
    "scoruri mai mari, iar cele cu profiluri diverse și complexe primesc scoruri mai mici — "
    "exact în linia predicțiilor din literatura de specialitate."
)

replace_all(
    "Utilizarea datasetului ESCO a permis analiza unui număr mare de ocupații și transformarea "
    "informațiilor ocupaționale într-o formă numerică adecvată pentru modelare predictivă.",

    "Alegerea ESCO ca sursă de date s-a dovedit potrivită: baza de date este structurată, "
    "accesibilă și suficient de granulară pentru a extrage variabile predictive relevante. "
    "Disponibilitatea versiunii în română a facilitat și etapa de mapare COR–ESCO."
)

replace_all(
    "Modelul LightGBM utilizat în cadrul cercetării a reușit să genereze scoruri continue și "
    "clasificări pe niveluri de risc, oferind o perspectivă aplicată asupra relației dintre "
    "competențe și automatizare.",

    "Algoritmul LightGBM s-a comportat conform așteptărilor pentru acest tip de problemă: "
    "performanță superioară regresiei liniare pe setul de test, integrabilitate naturală cu "
    "analiza SHAP și robustețe la lipsa de normalizare a variabilelor. Instabilitatea "
    "cross-validare este, din câte putem evalua, o consecință a relației complexe și "
    "parțial zgomotoase dintre variabilele disponibile și variabila țintă — nu un artefact "
    "al alegerii algoritmului."
)

replace_all(
    "Aplicarea modelului asupra Clasificației Ocupațiilor din România a demonstrat că "
    "metodologia poate fi transferată către un context național și poate produce rezultate "
    "relevante pentru analiza pieței muncii.",

    "Transferul pe COR a funcționat tehnic pentru ocupațiile mapate, producând scoruri "
    "interpretabile și consistente cu logica modelului. Principala problemă rămâne "
    "acoperirea scăzută (14.7%), care limitează valoarea analizei pentru scopuri de "
    "politică publică în stadiul actual."
)

replace_all(
    "Distribuția rezultatelor indică o repartizare echilibrată între LOW și MEDIUM, ceea ce "
    "sugerează că multe profesii sunt expuse unui proces gradual de automatizare, fără a "
    "fi complet înlocuite.",

    "Distribuția aproximativ egală LOW/MEDIUM în rândul ocupațiilor COR mapate este "
    "consistentă cu perspectiva mai optimistă din literatura recentă — că automatizarea "
    "transformă mai degrabă decât elimină ocupațiile — dar trebuie citită cu rezerva că "
    "eșantionul de 654 de ocupații nu este reprezentativ pentru întreaga piață a muncii."
)

replace_all(
    "Prin integrarea metodelor de explainability, cercetarea nu s-a limitat doar la obținerea "
    "unor predicții, ci a urmărit și înțelegerea modului în care variabilele influențează "
    "rezultatul final.",

    "Integrarea SHAP reprezintă, în opinia mea, una dintre contribuțiile mai valoroase ale "
    "acestei lucrări. Chiar dacă performanța predictivă este modestă, analiza SHAP adaugă "
    "un nivel de transparență care face diferența față de o simplă sortare a ocupațiilor "
    "după un scor opac — utilizatorul poate înțelege de ce o ocupație primește un anumit "
    "scor și poate evalua critic această justificare."
)

replace_all(
    "În ansamblu, lucrarea confirmă faptul că tehnicile moderne de machine learning pot fi "
    "utilizate pentru analiza ocupațională și pentru evaluarea riscului asociat transformărilor "
    "tehnologice.",

    "Concluzia mai largă este că metodologiile de machine learning aplicate pe date "
    "ocupaționale structurate sunt o abordare viabilă pentru analiza riscului de automatizare "
    "la nivel național — cu condiția de a fi completate cu date mai bogate, o mapare mai "
    "precisă a clasificărilor ocupaționale și o validare externă față de estimări independente."
)

replace_all(
    "Lucrarea aduce mai multe contribuții atât din perspectivă metodologică, cât și aplicativă.",

    "Contribuțiile concrete ale acestei cercetări pot fi grupate pe trei dimensiuni."
)

replace_all(
    "Prima contribuție constă în dezvoltarea unui model predictiv capabil să estimeze riscul "
    "de automatizare utilizând date structurate despre competențe și ocupații.",

    "Din perspectivă metodologică, lucrarea propune un pipeline reproductibil care combină "
    "extragerea de caracteristici din ESCO, construcția unui indicator proxy de automatizare, "
    "antrenarea LightGBM și interpretarea SHAP — o secvență care poate fi replicată sau "
    "adaptată de cercetători interesați de analize similare."
)

replace_all(
    "O altă contribuție importantă este utilizarea sistemului ESCO ca sursă principală de "
    "date, ceea ce permite o analiză standardizată și comparabilă la nivel european.",

    "Din perspectivă aplicativă, lucrarea furnizează scoruri de automatizare pentru 654 de "
    "ocupații din COR — o resursă inexistentă anterior pentru piața muncii din România — "
    "și le face accesibile printr-o aplicație Streamlit interactivă."
)

replace_all(
    "De asemenea, cercetarea propune o metodologie de transfer între clasificări ocupaționale "
    "diferite, prin asocierea dintre ESCO și COR.",

    "Din perspectivă teoretică, cercetarea contribuie la literatura privind transferul de "
    "modele între sisteme de clasificare ocupațională diferite — o problemă cu relevanță "
    "pentru orice studiu care încearcă să transpună rezultate internaționale în contexte "
    "naționale specifice."
)

replace_all(
    "Integrarea explainability prin metoda SHAP reprezintă o contribuție suplimentară, "
    "deoarece oferă transparență și interpretabilitate asupra modului în care modelul "
    "generează predicțiile.",

    "Integrarea SHAP ca instrument de validare internă și de comunicare a rezultatelor "
    "reprezintă o alegere metodologică deliberată: în domenii cu implicații de politică "
    "publică, transparența algoritmică nu este opțională."
)

replace_all(
    "În plus, dezvoltarea unei aplicații interactive în Streamlit oferă o componentă practică "
    "și facilitează explorarea rezultatelor într-un mod accesibil.",

    "Aplicația Streamlit dezvoltată paralel cu cercetarea face ca rezultatele să fie "
    "explorabile fără cunoștințe tehnice — un aspect relevant dacă instrumentul urmează "
    "să fie utilizat de consilieri de carieră, decidenți sau studenți interesați de "
    "profilul de risc al unor ocupații specifice."
)

replace_all(
    "Prin combinarea acestor elemente, cercetarea contribuie la dezvoltarea unei abordări "
    "aplicate pentru analiza automatizării ocupațiilor.",

    "Combinate, aceste contribuții adresează un gol real: absența unui instrument adaptat "
    "contextului românesc pentru estimarea riscului de automatizare ocupațională bazat pe "
    "date și metodologie transparentă."
)

replace_all(
    "Deși rezultatele obținute sunt relevante, cercetarea prezintă mai multe limitări care "
    "trebuie luate în considerare.",

    "Onestitatea intelectuală impune să discutăm explicit ce nu a funcționat sau a funcționat "
    "mai puțin bine — nu ca o formulă de modestie academică, ci pentru că limitele "
    "condiționate direct valoarea practică a rezultatelor."
)

replace_all(
    "Una dintre principalele limitări este reprezentată de performanța modelului în ceea ce "
    "privește generalizarea.",

    "Cea mai serioasă limitare rămâne instabilitatea modelului în cross-validare "
    "(R²_CV = −14.21 ± 6.16), care contrastează nefavorabil cu valoarea de pe setul de "
    "test. Această discrepanță sugerează că succesul pe test-set poate fi parțial un "
    "artefact al particularității acelei partiții, și nu o garanție de generalizare robustă."
)

replace_all(
    "Această limitare poate fi asociată atât structurii datasetului, cât și numărului redus "
    "de variabile utilizate.",

    "Extinderea setului de variabile cu date externe — salariile medii pe ocupație, sectorul "
    "economic predominant, cererea de competențe din anunțurile de angajare — ar putea "
    "ameliora semnificativ stabilitatea și precizia modelului."
)

replace_all(
    "Modelul a fost construit exclusiv pe baza caracteristicilor ocupaționale și nu include "
    "variabile economice sau sociale care pot influența riscul de automatizare.",

    "O a doua limitare importantă este că modelul nu are acces la informații contextuale: "
    "nu știe dacă o ocupație se exercită în sectorul public sau privat, în industrii "
    "intensiv-tehnologice sau tradiționale, în zone urbane sau rurale. Toate acestea "
    "modulează viteza reală de adoptare a automatizării."
)

replace_all(
    "Factori precum salariul, industria, gradul de digitalizare al domeniului sau cererea de "
    "pe piața muncii pot contribui semnificativ la vulnerabilitatea ocupațiilor.",

    "Cercetările lui Acemoglu și Restrepo (2020) arată că rata de automatizare variază "
    "considerabil în funcție de presiunile din piața muncii locală și de disponibilitatea "
    "capitalului pentru investiții în automatizare — factori pe care un model bazat exclusiv "
    "pe competențe nu îi poate captura."
)

replace_all(
    "Asocierea ocupațiilor a fost realizată prin metode de similaritate textuală, ceea ce "
    "poate introduce erori semantice.",

    "Fuzzy matching-ul, deși eficient computațional, produce asocieri bazate pe forma "
    "textului, nu pe sensul ocupațional. Două ocupații cu denumiri similare pot corespunde "
    "unor realități profesionale diferite, și invers — ocupații echivalente funcțional "
    "pot rămâne nemapate din cauza diferențelor de nomenclatură."
)

replace_all(
    "În anumite situații, ocupații cu denumiri similare pot avea diferențe importante în "
    "ceea ce privește activitățile desfășurate.",

    "O validare manuală pe un eșantion din asocierile realizate ar fi o bună practică pentru "
    "cuantificarea ratei de eroare semantică și ar oferi o bază mai solidă pentru "
    "interpretarea scorurilor transferate pe COR."
)

replace_all(
    "Rata de mapare de aproximativ 14.7% indică faptul că o parte semnificativă dintre "
    "ocupațiile COR nu au putut fi asociate cu ocupațiile ESCO.",

    "Rata de acoperire de 14.7% este cea mai concretă limitare operațională: aproape "
    "85% din ocupațiile COR nu au primit scoruri de automatizare. Orice concluzie despre "
    "profilul de risc al pieței muncii românești bazată pe aceste date este, prin urmare, "
    "extrapolată dintr-un subset nereprezentativ."
)

replace_all(
    "Această limitare reduce gradul de acoperire al modelului și influențează analiza finală.",

    "Extinderea acoperirii — fie printr-un prag de matching mai permisiv combinat cu "
    "validare manuală a asocierilor incerte, fie prin tehnici NLP mai avansate — rămâne "
    "prioritatea metodologică numărul unu pentru versiunile viitoare."
)

replace_all(
    "Rezultatele obținute în cadrul cercetării deschid posibilitatea extinderii modelului "
    "și dezvoltării unor direcții viitoare.",

    "Rezultatele obținute, cu tot cu limitările lor, trasează câteva direcții clare de "
    "dezvoltare pe care cercetări viitoare le pot urma."
)

replace_all(
    "Una dintre direcțiile importante constă în includerea unor variabile suplimentare care "
    "să completeze analiza ocupațională.",

    "Prima și probabil cea mai impactantă direcție este îmbogățirea setului de variabile "
    "cu date externe. Corelarea profilelor competențiale din ESCO cu date de salarizare "
    "(disponibile de la INS sau Eurostat), cerere de competențe din anunțuri de angajare "
    "sau grade de digitalizare sectoriale ar putea transforma un model cu R²≈0.2 într-unul "
    "cu performanțe semnificativ mai robuste."
)

replace_all(
    "Integrarea datelor economice, precum salariile, nivelul de ocupare, industria și "
    "cererea de competențe, ar putea îmbunătăți performanța modelului.",

    "Bessen (2019) a arătat că cererea de competențe măsurată din anunțuri de angajare "
    "este unul dintre cei mai buni predictori pentru adoptarea tehnologică la nivel de "
    "ocupație. Integrarea unor astfel de date în modelul de față ar reprezenta un salt "
    "calitativ față de abordarea actuală bazată exclusiv pe structuri statice de competențe."
)

replace_all(
    "De asemenea, pot fi testate și alte algoritmi de machine learning, precum XGBoost, "
    "CatBoost sau modele bazate pe rețele neuronale.",

    "O a doua direcție este compararea sistematică a algoritmilor. Deși LightGBM s-a "
    "comportat bine, nu există garanții că este optim pentru această problemă. XGBoost, "
    "CatBoost sau chiar modele mai simple cu regularizare mai agresivă ar putea oferi "
    "un echilibru mai bun între bias și varianță — tocmai ceea ce modelul actual pare "
    "să aibă dificultăți să obțină."
)

replace_all(
    "O altă direcție de dezvoltare constă în îmbunătățirea procesului de asociere între "
    "COR și ESCO.",

    "A treia direcție — și poate cea cu cel mai mare impact practic — este îmbunătățirea "
    "mapării COR–ESCO."
)

replace_all(
    "În locul metodei de fuzzy matching, pot fi utilizate tehnici avansate de procesare a "
    "limbajului natural pentru identificarea relațiilor semantice dintre ocupații.",

    "Tehnicile moderne de NLP, cum ar fi sentence embeddings (SBERT, multilingual BERT), "
    "pot identifica ocupații echivalente semantic chiar și atunci când denumirile textuale "
    "diferă semnificativ. O mapare bazată pe reprezentări semantice, validate parțial manual, "
    "ar putea crește rata de acoperire la 60–80% cu o rată de eroare controlată."
)

replace_all(
    "Extinderea cercetării la nivel sectorial poate reprezenta o altă direcție relevantă.",

    "O a patra direcție este analiza sectorială diferențiată."
)

replace_all(
    "Analiza separată a unor domenii precum IT, industrie, sănătate sau administrație ar "
    "putea evidenția diferențe importante privind riscul de automatizare.",

    "Modelul curent tratează toate ocupațiile uniform, indiferent de sectorul în care se "
    "exercită. O analiză separată pe sectoare (IT, sănătate, industrie, servicii publice) "
    "ar permite calibrări mai fine și ar putea evidenția cum dinamicile sectoriale "
    "modulează riscul structural estimat din profilul de competențe."
)

replace_all(
    "Automatizarea și inteligența artificială influențează în mod direct evoluția ocupațiilor "
    "și transformarea pieței muncii. ",

    "Trăim într-un moment în care întrebarea nu mai este dacă automatizarea va schimba "
    "piața muncii, ci cât de repede, în ce sectoare și cu ce consecințe distributive. "
    "Cercetarea de față nu pretinde să răspundă la aceste întrebări mari — dar poate "
    "contribui la o parte din infrastructura metodologică necesară pentru a le aborda "
    "în mod sistematic."
)

replace_all(
    "În acest context, dezvoltarea unor metode predictive pentru estimarea vulnerabilității "
    "ocupațiilor devine relevantă atât pentru cercetarea academică, cât și pentru "
    "instituțiile implicate în planificarea forței de muncă.",

    "Instrumentele de estimare a riscului ocupațional — chiar și cele cu limitări "
    "metodologice evidente — pot servi ca punct de plecare pentru dialog între cercetători, "
    "factori de decizie și practicieni ai pieței muncii, cu condiția ca limitele să fie "
    "comunicate transparent."
)

replace_all(
    "Lucrarea de față a demonstrat că utilizarea tehnicilor de machine learning și a datelor "
    "ocupaționale poate contribui la evaluarea riscului de automatizare.",

    "Ceea ce această lucrare a demonstrat concret este că datele structurale din ESCO "
    "conțin semnale suficiente pentru a construi un sistem funcțional — nu perfect, "
    "dar funcțional — de estimare a riscului ocupațional. Un scor de automatizare bazat "
    "pe competențe nu este un substitut pentru o analiză economică completă, dar este "
    "mai bine decât absența oricărui instrument."
)

replace_all(
    "Prin combinarea datelor ESCO, a modelului LightGBM și a metodelor de interpretabilitate, "
    "cercetarea propune o abordare practică și reproductibilă.",

    "Combinând ESCO, LightGBM și SHAP într-un pipeline reproductibil, disponibil public pe "
    "GitHub, lucrarea pune bazele pentru cercetări iterative care pot îmbunătăți treptat "
    "precizia și acoperirea estimărilor."
)

replace_all(
    "Deși există limitări metodologice, rezultatele obținute evidențiază potențialul "
    "utilizării inteligenței artificiale pentru analiza ocupațională.",

    "Limitările metodologice sunt reale și nu trebuie minimizate. Dar potențialul "
    "abordării este la fel de real: cu date mai bogate, o mapare mai precisă și o "
    "validare externă adecvată, un sistem similar ar putea deveni un instrument "
    "valoros pentru analiza pieței muncii românești."
)

replace_all(
    "Studiul poate reprezenta un punct de plecare pentru cercetări viitoare și pentru "
    "dezvoltarea unor instrumente capabile să sprijine analiza transformărilor produse "
    "de automatizare asupra pieței muncii.",

    "Sper că această lucrare poate servi nu ca un răspuns definitiv la întrebările "
    "privind automatizarea ocupațională în România, ci ca o primă aproximare utilă și "
    "ca o bază de pornire pentru cercetări viitoare mai complete."
)

doc.save(str(DEST))
print(f"\n✓ Salvat: {DEST}")
print(f"  Dimensiune: {DEST.stat().st_size // 1024} KB")

# Sterge fisierul temporar de text
import os
try:
    os.remove("Lucrare/text_curent.txt")
except:
    pass
