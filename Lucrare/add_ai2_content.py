"""
add_ai2_content.py — Adaugă conținut AI 2 în Estimarea probabilitatii_APA - uman.docx
Produce: Estimarea probabilitatii_APA - complet.docx
"""
import sys, shutil
from pathlib import Path
from copy import deepcopy
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

sys.stdout.reconfigure(encoding="utf-8")

SRC  = Path("Lucrare/Estimarea probabilitatii_APA - uman.docx")
DEST = Path("Lucrare/Estimarea probabilitatii_APA - complet.docx")
shutil.copy2(SRC, DEST)
doc = Document(str(DEST))


# ── Utilitare ─────────────────────────────────────────────────────────────────

def find_para(fragment):
    """Returnează primul paragraf care conține fragmentul dat."""
    for p in doc.paragraphs:
        if fragment in p.text:
            return p
    return None


def _clone_normal_elem(ref_para):
    """Clonează elementul XML al unui paragraf 'normal', șterge run-urile."""
    elem = deepcopy(ref_para._element)
    for child in list(elem):
        if child.tag != qn("w:pPr"):
            elem.remove(child)
    return elem


def _make_elem_with_text(ref_para, text, bold=False):
    elem = _clone_normal_elem(ref_para)
    r = OxmlElement("w:r")
    if bold:
        rPr = OxmlElement("w:rPr")
        b = OxmlElement("w:b")
        rPr.append(b)
        r.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    r.append(t)
    elem.append(r)
    return elem


def _make_heading_elem(doc, text, level=2):
    elem = OxmlElement("w:p")
    pPr  = OxmlElement("w:pPr")
    pStyle = OxmlElement("w:pStyle")
    style_id = None
    for s in doc.styles:
        if s.name.lower() == f"heading {level}":
            style_id = s.style_id
            break
    if style_id:
        pStyle.set(qn("w:val"), style_id)
        pPr.append(pStyle)
    elem.append(pPr)
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = text
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    r.append(t)
    elem.append(r)
    return elem


def insert_block(anchor_fragment, blocks, ref_fragment=None):
    """
    Inserează o listă de blocuri după ancora dată.
    blocks = list de (text, tip) unde tip ∈ {"body", "h2", "h3"}
    ref_fragment: un paragraf normal pentru a-i copia stilul (opțional)
    """
    anchor = find_para(anchor_fragment)
    if anchor is None:
        print(f"  !! SKIP — ancora nu a fost găsită: '{anchor_fragment[:60]}'")
        return
    ref = find_para(ref_fragment) if ref_fragment else anchor
    if ref is None:
        ref = anchor

    current = anchor._element
    for tip, text in blocks:
        if tip == "body":
            new_elem = _make_elem_with_text(ref, text)
        elif tip == "bold":
            new_elem = _make_elem_with_text(ref, text, bold=True)
        elif tip == "h2":
            new_elem = _make_heading_elem(doc, text, level=2)
        elif tip == "h3":
            new_elem = _make_heading_elem(doc, text, level=3)
        else:
            new_elem = _make_elem_with_text(ref, text)
        current.addnext(new_elem)
        current = new_elem
    print(f"  ✓ Inserat {len(blocks)} blocuri după: '{anchor_fragment[:55]}'")


# ═══════════════════════════════════════════════════════════════════════════════
# 1. ABSTRACT — mențiune scurtă AI 2
# ═══════════════════════════════════════════════════════════════════════════════
print("\n=== ABSTRACT — adăugare AI 2 ===")
insert_block(
    "constituind un punct de plecare pentru analize mai fine ale pieței muncii românești.",
    [
        ("body",
         "Lucrarea include și o componentă experimentală de evaluare aprofundată — denumită AI 2 — "
         "bazată pe modelul de limbaj GPT-4o (OpenAI), integrată printr-un mecanism hibrid care "
         "combină un scoring determinist derivat din profilul de competențe cu estimările semantice "
         "ale modelului de limbaj. Această componentă introduce o distincție conceptuală față de "
         "evaluarea tradițională: separarea riscului de înlocuire completă (replacementRisk) de "
         "intensitatea impactului AI asupra activităților curente (aiImpactScore). Un profesor, de "
         "pildă, poate înregistra replacementRisk scăzut (8/100) — ocupația nu dispare — și "
         "aiImpactScore ridicat (65/100), reflectând transformarea substanțială a modului de lucru "
         "prin instrumente inteligente. Componenta AI 2 integrează un sistem de colectare automată "
         "din surse externe — arXiv, Hugging Face Hub, OpenAlex, bloguri oficiale — cu stocare "
         "locală în SQLite și ranking compozit bazat pe relevanță, recență și credibilitate.",),
    ],
    ref_fragment="constituind un punct de plecare"
)

# ═══════════════════════════════════════════════════════════════════════════════
# 2. CAP 3 — Secțiunile 3.9 și 3.10 (după workflow-ul existent)
# ═══════════════════════════════════════════════════════════════════════════════
print("\n=== CAP 3 — Secțiunile 3.9 și 3.10 ===")
insert_block(
    "Pentru o reprezentare vizuală a fluxului de lucru, poate fi introdusă o schemă metodologică",
    [
        # ── 3.9 ──
        ("h2",
         "3.9 Componenta AI 2 — Evaluator Hibrid cu Model de Limbaj de Mari Dimensiuni"),

        ("body",
         "Paralel cu modelul de regresie LightGBM, aplicația include o a doua componentă de "
         "evaluare — AI 2 — care urmărește să compenseze câteva dintre limitele inerente ale "
         "primului model. Cel mai evident neajuns al AI 1 este că lucrează exclusiv cu șase "
         "variabile structurale din ESCO, fără să poată incorpora informații despre evoluțiile "
         "recente din domeniu sau să distingă tipuri calitativ diferite de impact al automatizării. "
         "AI 2 încearcă să acopere parțial această lipsă printr-un mecanism hibrid care combină "
         "un scoring determinist cu evaluarea unui model de limbaj de mari dimensiuni."),

        ("body",
         "Prima etapă a mecanismului calculează un scor determinist pornind de la profilul de "
         "competențe al ocupației analizate. Scorul inițial furnizat de AI 1 este corectat prin "
         "formula: det_score = scor_inițial × 0,40 + skill_based × 0,60, unde skill_based "
         "reflectă proporția competențelor cu risc ridicat față de total. Competențele sunt "
         "clasificate în trei categorii — risc înalt, mediu și redus — pe baza unui lexicon de "
         "termeni asociați sarcinilor rutinizabile (procesare, înregistrare, calcul, sortare) "
         "față de termeni care indică sarcini nerutinizabile (coordonare, evaluare, consiliere, "
         "creație, adaptare). Scorul final combină în proporții egale componentele deterministă "
         "și GPT: scor_final = round(det_score × 0,5 + scor_GPT × 0,5)."),

        ("body",
         "A doua etapă trimite modelului GPT-4o un prompt structurat care conține profilul de "
         "competențe al ocupației, scorul inițial AI 1 și o selecție de cel mult cinci surse "
         "externe relevante extrase din cache-ul local. Modelul este instruit să returneze un "
         "obiect JSON cu câmpuri predefinite: scor GPT, riscul de înlocuire (replacementRisk, "
         "0–100), impactul AI (aiImpactScore, 0–100), modificarea față de scorul inițial "
         "(scoreChange), motivul ajustării (scoreAdjustmentReason), recomandări profesionale, "
         "competențe transferabile, ocupații alternative și o listă de progrese tehnologice "
         "relevante. Câmpul scoreChange este recalculat întotdeauna determinist — ca diferență "
         "între scorul final și scorul inițial — pentru a evita inconsistențele generate de "
         "variabilitatea modelului de limbaj."),

        ("body",
         "O contribuție conceptuală importantă a componentei AI 2 este introducerea explicită "
         "a distincției dintre replacementRisk și aiImpactScore. Cei doi indicatori capturează "
         "dimensiuni calitativ diferite ale impactului automatizării, care în literatura de "
         "specialitate sunt adesea tratate împreună, dar care în practică pot evolua independent. "
         "replacementRisk estimează probabilitatea că ocupația va dispărea complet din piața "
         "muncii ca urmare a automatizării. aiImpactScore capturează gradul în care instrumentele "
         "AI transformă activitățile zilnice ale practicianului, chiar și fără a elimina ocupația. "
         "Un profesor poate înregistra aiImpactScore de 65/100 pentru că platformele adaptive, "
         "generatoarele de materiale didactice și sistemele de evaluare automată îi schimbă "
         "substanțial rutina, fără a-l face redundant. Invers, un operator de prelucrare a datelor "
         "poate avea ambele scoruri ridicate, indicând o probabilitate mai mare de înlocuire totală."),

        ("body",
         "Distincția este relevantă practic: politicile de reconversie profesională vizează diferit "
         "o ocupație cu replacementRisk ridicat față de una cu aiImpactScore ridicat. Prima necesită "
         "redirecționare spre domenii noi; a doua necesită upskilling în utilizarea instrumentelor "
         "AI. Confundarea celor două dimensiuni conduce la intervenții de politici publice inadecvate, "
         "ceea ce justifică includerea lor separată în instrumentul de față."),

        ("body",
         "Pentru a reduce variabilitatea răspunsurilor GPT-4o și costurile de API, aplicația "
         "include un mod demonstrativ cu rezultate pre-calculate pentru cinci ocupații "
         "reprezentative: Contabil, Designer grafic, Asistent medical, Profesor și Analist "
         "programator. Rezultatele demonstrative acoperă diversitatea profilelor de risc și "
         "permit testarea interfeței fără apeluri API. Căutarea ocupației se face prin "
         "potrivire fuzzy pe lista de chei demonstrative, cu fallback la evaluarea live."),

        # ── 3.10 ──
        ("h2",
         "3.10 Sistemul de Colectare și Gestionare a Surselor Externe"),

        ("body",
         "Pentru ca evaluarea AI 2 să nu se bazeze exclusiv pe cunoștințele parametrice ale "
         "modelului GPT-4o — cunoștințe cu dată de expirare și potențial inconsistente cu "
         "evoluțiile recente — am integrat un sistem modular de colectare și gestionare a "
         "surselor externe. Arhitectura este de tip cache-first: colectarea se face periodic "
         "sau la cerere, nu la fiecare evaluare, eliminând latența căutărilor live."),

        ("body",
         "Sistemul include cinci tipuri de colectori. Colectorul arXiv extrage articole recente "
         "din categoriile cs.AI, cs.LG, cs.CL, cs.CV, cs.RO și econ.GN, fără autentificare — "
         "API-ul arXiv este public. Maturitatea fiecărui articol se inferează din vechimea "
         "publicației: sub 180 de zile înseamnă 'emerging', sub doi ani 'growing', altfel 'high'. "
         "Colectorul Hugging Face Hub interogează modelele publice cu cel puțin 10.000 de "
         "descărcări din zece tipuri de task-uri: generare text, recunoaștere vocală, clasificare "
         "imagini, detectare obiecte, traducere, sumarizare, clasificare text, question answering, "
         "NER și image-to-text. Impactul estimat al unui model se derivă din numărul de aprecieri "
         "(likes ≥ 1.000 = înalt, ≥ 100 = mediu, altfel = redus)."),

        ("body",
         "Colectorul OpenAlex exploatează API-ul gratuit cu parametru polite_pool (adresă de "
         "e-mail) care îmbunătățește limita de rată. Se filtrează articole de tip 'article' "
         "publicate după un an de referință și se sortează descrescător după numărul de citări. "
         "Colectorul RSS/bloguri oficiale parsează fluxurile de știri de la OpenAI, Google "
         "DeepMind, Meta AI, Anthropic, Hugging Face Blog și MIT Technology Review, suportând "
         "atât RSS 2.0 cât și Atom. Colectorul Discord este opțional — necesită token de bot — "
         "și stochează exclusiv date non-personale: URL-ul mesajului, primul rând non-gol ca "
         "titlu, rezumatul (maxim 500 de caractere) și denumirea canalului. Această componentă "
         "este etichetată explicit ca 'semnal curatorial', nu sursă academică."),

        ("body",
         "Sursele colectate sunt stocate într-o bază de date SQLite locală cu mod WAL "
         "(Write-Ahead Logging), pentru citiri concurente în mediul multi-thread al Streamlit. "
         "Deduplicarea se realizează prin similaritate Jaccard pe seturile de cuvinte din titlu, "
         "cu prag > 0,85. Un modul de ranking compozit atribuie fiecărei surse un scor calculat "
         "astfel: scor_compus = (relevanță × 0,50 + recență × 0,20 + impact × 0,20 + dovezi × "
         "0,10) × credibilitate_sursă. Credibilitatea reflectă tipul sursei: OpenAlex (0,90) > "
         "arXiv (0,85) > bloguri oficiale (0,80) > Hugging Face (0,75) > estimări GPT (0,55) > "
         "Discord (0,40). Ierarhia operaționalizează diferența dintre surse peer-reviewed și "
         "surse de semnal curatorial."),

        ("body",
         "Dacă cache-ul este gol sau nu conține surse relevante pentru ocupația analizată, "
         "componenta AI 2 revine la estimările proprii ale GPT-4o, fără a afecta funcționalitatea "
         "de bază. Interfața aplicației Streamlit include un buton de actualizare în bara laterală "
         "care declanșează colectarea pentru toate sursele disponibile și un rezumat al stării "
         "cache-ului: număr total de înregistrări, distribuție pe tip de sursă și data ultimei "
         "actualizări. Această transparență față de utilizator este importantă pentru evaluarea "
         "calității surselor care au influențat o anumită estimare."),
    ],
    ref_fragment="Procesul metodologic utilizat în cadrul cercetării"
)

# ═══════════════════════════════════════════════════════════════════════════════
# 3. CAP 4 — Secțiunea 4.11 cu rezultatele AI 2
# ═══════════════════════════════════════════════════════════════════════════════
print("\n=== CAP 4 — Secțiunea 4.11 ===")
insert_block(
    "În ansamblu, cercetarea demonstrează că datele ESCO conțin semnale suficiente pentru",
    [
        ("h2",
         "4.11 Rezultatele Componentei AI 2 — Evaluator Hibrid"),

        ("body",
         "Componenta AI 2 a fost testată pe cinci ocupații reprezentative, alese pentru a "
         "ilustra atât diversitatea profilelor de risc, cât și distincția conceptuală dintre "
         "replacementRisk și aiImpactScore. Prezentăm în continuare rezultatele obținute "
         "pentru fiecare, cu observații privind coerența cu estimările AI 1 și cu literatura "
         "de specialitate."),

        ("body",
         "Contabilul (cod ISCO 2411) obține un scor final AI 2 de 61/100, față de 56/100 "
         "estimat de AI 1. Diferența modestă reflectă că AI 2 identifică un risc de "
         "rutinizare mai ridicat în sarcinile de reconciliere și raportare, compensat parțial "
         "de capacitatea de interpretare contextuală. replacementRisk este de 38/100 — ocupația "
         "este vulnerabilă în segmentele sale mai repetitive, dar nu complet substituibilă — "
         "iar aiImpactScore atinge 72/100, reflectând transformarea semnificativă a fluxului "
         "de lucru prin instrumente de reconciliere automată, audit asistat de AI și generare "
         "de rapoarte financiare. Profesia nu dispare, dar se reconfigurează substanțial în "
         "jurul sarcinilor de interpretare și decizie."),

        ("body",
         "Asistentul medical (ISCO 3221) are scor final de 33/100, cu replacementRisk de "
         "12/100 și aiImpactScore de 45/100. Profilul este consistent cu literatura de "
         "specialitate: componenta de îngrijire directă, evaluare clinică și comunicare "
         "empatică plasează ocupația în zona de risc scăzut, deși instrumentele de "
         "monitorizare a parametrilor vitali și de suport la decizie clinică modifică parțial "
         "fluxul de lucru. Este notabil că aiImpactScore rămâne moderat chiar și pentru o "
         "ocupație cu replacementRisk scăzut — dovadă că transformarea nu înseamnă neapărat "
         "risc de eliminare."),

        ("body",
         "Cel mai ilustrativ exemplu pentru distincția conceptuală introdusă de componenta "
         "AI 2 este Profesorul (ISCO 2320): replacementRisk de 8/100 se combină cu "
         "aiImpactScore de 65/100. Ocupația nu va dispărea — profesorul rămâne indispensabil "
         "pentru formarea critică, mentoring și dinamica socializării — dar instrumentele de "
         "tutoring adaptiv, generare automată de conținut și evaluare asistată de AI "
         "transformă substanțial modul în care predarea se desfășoară. Aceasta este o "
         "distincție pe care AI 1, cu un scor unic de 24/100 (crescut la 28/100 de AI 2), "
         "nu o putea exprima. Profesorul viitorului va orchestra instrumente inteligente, "
         "nu va fi înlocuit de ele."),

        ("body",
         "Analistul programator (ISCO 2512) înregistrează cel mai ridicat aiImpactScore din "
         "setul de testare: 80/100, cu replacementRisk de 22/100 și scor final de 48/100. "
         "Combinația reflectă realitatea documentată în industria software: asistenți de cod "
         "(GitHub Copilot, Cursor), generatoare de teste și sisteme de review automat "
         "transformă radical fluxul de lucru, dar programatorul care poate orchestra aceste "
         "instrumente, valida output-ul și lua decizii arhitecturale rămâne mai valoros, "
         "nu mai puțin relevant. Scorul relativ modest de 48/100 față de aiImpactScore "
         "ridicat ilustrează că impactul AI nu echivalează cu riscul de înlocuire."),

        ("body",
         "Designerul grafic (ISCO 2166) prezintă un profil intermediar: scor final de "
         "49/100, replacementRisk de 28/100, aiImpactScore de 75/100. Generatoarele de "
         "imagini (Midjourney, DALL-E, Stable Diffusion) și platformele de design asistat "
         "AI transformă radical fluxul creativ, dar judecata estetică, înțelegerea brandului "
         "clientului și comunicarea cu stakeholderii rămân capacități eminamente umane. "
         "Designerul se transformă dintr-un executor tehnic într-un curator și orchestrator "
         "de instrumente AI — un rol care necesită competențe noi, nu dispariția profesiei."),

        ("body",
         "Distribuția celor cinci exemple confirmă că scorurile AI 2 sunt consistente cu "
         "estimările AI 1 la nivel de direcție — corelația pozitivă dintre cele două seturi "
         "de scoruri finale este clară — dar aduc o granularitate suplimentară prin separarea "
         "dimensiunilor de risc și prin integrarea informațiilor contextuale despre evoluțiile "
         "tehnologice recente. Componenta AI 2 produce, de regulă, scoruri ușor mai ridicate "
         "decât AI 1 pentru ocupațiile cu profil mixt, reflectând sensibilitatea mai mare față "
         "de evoluțiile recente din domeniu."),
    ],
    ref_fragment="Analiza SHAP a validat coerența modelului: variabilele proiectate"
)

# ═══════════════════════════════════════════════════════════════════════════════
# 4. CAP 5.2 — Contribuții: adăugare AI 2
# ═══════════════════════════════════════════════════════════════════════════════
print("\n=== CAP 5.2 — Contribuții AI 2 ===")
insert_block(
    "Combinate, aceste contribuții adresează un gol real: absența unui instrument adaptat",
    [
        ("body",
         "O a patra contribuție, metodologică și conceptuală, o reprezintă introducerea "
         "componentei AI 2 și a distincției explicite dintre replacementRisk și aiImpactScore. "
         "Cele două dimensiuni operaționalizează o nuanță din literatura de specialitate — că "
         "automatizarea transformă ocupațiile în moduri calitativ diferite de la o ocupație la "
         "alta — și fac ca instrumentul să producă mai multă informație utilă decât un scor "
         "unic de risc. Componenta AI 2 demonstrează, de asemenea, că integrarea modelelor de "
         "limbaj de mari dimensiuni cu scoring determinist și surse externe verificabile poate "
         "produce evaluări mai bogate decât oricare dintre aceste surse în mod individual."),

        ("body",
         "Nu în ultimul rând, sistemul de colectare automată a surselor externe — arXiv, "
         "Hugging Face Hub, OpenAlex, RSS — reprezintă o contribuție practică la problema "
         "actualizării continue a cunoașterii în instrumente de evaluare a riscului. Cache-ul "
         "SQLite cu ranking compozit și deduplicare Jaccard poate fi reutilizat independent "
         "de contextul specific al acestei lucrări, ca modul de colectare a literaturii de "
         "specialitate relevante pentru orice domeniu al pieței muncii."),
    ],
    ref_fragment="absența unui instrument adaptat contextului românesc"
)

# ═══════════════════════════════════════════════════════════════════════════════
# 5. CAP 5.3 — Limitări: adăugare limitări AI 2
# ═══════════════════════════════════════════════════════════════════════════════
print("\n=== CAP 5.3 — Limitări AI 2 ===")
insert_block(
    "Extinderea acoperirii — fie printr-un prag de matching mai permisiv combinat cu",
    [
        ("body",
         "Componenta AI 2 introduce propriile limitări, distincte de cele ale modelului "
         "statistic. GPT-4o este un model probabilistic ale cărui răspunsuri nu sunt perfect "
         "reproductibile — rulând aceeași interogare de două ori se pot obține scoruri ușor "
         "diferite. Utilizarea sa implică costuri de API (OpenAI), ceea ce limitează "
         "scalabilitatea la evaluări individuale sau seturi mici de ocupații, nu la procesări "
         "în masă. Calitatea evaluărilor depinde parțial de conținutul cache-ului de surse "
         "externe: dacă acesta este gol sau vechi, componenta revine la estimări bazate "
         "exclusiv pe cunoștințele parametrice ale modelului, fără context factual recent."),

        ("body",
         "Nu în ultimul rând, modelul GPT-4o poate manifesta biasuri proprii față de anumite "
         "ocupații sau domenii, moștenite din datele de pre-antrenare. Aceste biasuri sunt "
         "greu de detectat și cuantificat fără un set de referință validat extern. Mecanismul "
         "hibrid atenueaz parțial această problemă prin ancorarea evaluării GPT într-un scor "
         "determinist calculat din date obiective, dar nu o elimină complet. Utilizatorul "
         "aplicației trebuie să trateze rezultatele AI 2 ca estimări orientative, nu ca "
         "predicții exacte sau judecăți definitive despre viitorul unei ocupații."),
    ],
    ref_fragment="prioritatea metodologică numărul unu pentru versiunile viitoare."
)

# ═══════════════════════════════════════════════════════════════════════════════
# 6. CAP 5.4 — Direcții viitoare: AI 2
# ═══════════════════════════════════════════════════════════════════════════════
print("\n=== CAP 5.4 — Direcții viitoare AI 2 ===")
insert_block(
    "Modelul curent tratează toate ocupațiile uniform, indiferent de sectorul în care se",
    [
        ("body",
         "O a cincea direcție de dezvoltare privește componenta AI 2 și sistemul de surse "
         "externe. Pe termen scurt, clasificarea competențelor — de la lexiconul manual "
         "actual la un clasificator antrenat — ar putea crește precizia scorului determinist "
         "și reduce dependența de regulile hardcodate. Pe termen mediu, extinderea sistemului "
         "de colectare cu surse de cerere reală de competențe — anunțuri de angajare, date "
         "ANOFM — ar calibra estimările la dinamicile pieței muncii locale, nu doar la "
         "evoluțiile tehnologice globale."),

        ("body",
         "Pe termen lung, un model hibrid care integrează actualizări continue din surse "
         "externe — nu periodic, ci în timp real — ar putea deveni un instrument de "
         "monitorizare a pieței muncii cu aplicații practice directe pentru serviciile "
         "de orientare profesională și pentru instituțiile responsabile cu politicile de "
         "reconversie. Separarea replacementRisk de aiImpactScore, introdusă în această "
         "lucrare la nivel de instrument demonstrativ, ar putea fi validată empiric "
         "prin studii longitudinale care să urmărească evoluția reală a ocupațiilor în "
         "raport cu estimările generate la un moment dat."),
    ],
    ref_fragment="Modelul curent tratează toate ocupațiile uniform"
)

# ═══════════════════════════════════════════════════════════════════════════════
# Salvare
# ═══════════════════════════════════════════════════════════════════════════════
doc.save(str(DEST))
print(f"\n✓ Salvat: {DEST}")
print(f"  Dimensiune: {DEST.stat().st_size // 1024} KB")
print(f"  Paragrafe: {len(doc.paragraphs)}")
